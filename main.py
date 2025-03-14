# streamlit run "C:/Users/keno/OneDrive/Documents/Projects/DATA AUTOMIZER APP/app2.py"

import gspread
from oauth2client.service_account import ServiceAccountCredentials
from google.oauth2 import service_account
import smtplib
from email.mime.text import MIMEText
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sns
import numpy as np
import scipy.stats as stats
from wordcloud import WordCloud
import squarify
import plotly.graph_objects as go
import networkx as nx
from docx import Document
from io import BytesIO
import base64
import warnings
import io
import zipfile
import time
import google.generativeai as genai
import re 
import hashlib
import json
from PIL import Image
from statsmodels.tsa.seasonal import seasonal_decompose
from statsmodels.tsa.stattools import adfuller
import random

# STREAMLIT CSS AND CONFIGURATIONS ==============================================================================================================

logo = Image.open("insightfox_logo.jpeg")

st.set_page_config(
    page_title="InsightFox",
    page_icon=logo,
    layout="wide",  # or "wide" if you prefer
    initial_sidebar_state="auto"
)

st.set_option('client.showErrorDetails', False)

matplotlib.use('Agg')
warnings.filterwarnings("ignore", category=DeprecationWarning, message=".*pyplot.*")

st.markdown(
    """
    <style>
    section[data-testid="stMain"] > div[data-testid="stMainBlockContainer"] {
         padding-top: 0px;  # Remove padding completely
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Hide Streamlit style elements
hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                header {visibility: hidden;}
                </style>
                """
st.markdown(hide_st_style, unsafe_allow_html=True)

st.markdown("""
    <style>
    [data-testid="stTextArea"] {
        color: #FFFFFF;
    }
    </style>
    """, unsafe_allow_html=True)

# Set Montserrat font
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Montserrat', sans-serif;
    }
    </style>
    """, unsafe_allow_html=True)

# Change color of specific Streamlit elements
st.markdown("""
    <style>
    .st-emotion-cache-1o6s5t7 {
        color: #ababab !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stExpander {
        background-color: #F3F7EC;
        border-radius: 10px;
    }
    
    .stExpander > details {
        background-color: #F3F7EC;
        border-radius: 10px;
    }
    
    .stExpander > details > summary {
        background-color: #F3F7EC;
        border-radius: 10px 10px 0 0;
        padding: 10px;
    }
    
    .stExpander > details > div {
        background-color: #F3F7EC;
        border-radius: 0 0 10px 10px;
        padding: 10px;
    }
    
    .stCheckbox {
        background-color: #F3F7EC;
        border-radius: 5px;
        padding: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stButton > button {
        color: #F3F7EC;
        background-color: #bd5c34;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown(
    """
    <style>
    .streamlit-expanderHeader {
        font-size: 20px;
    }
    .streamlit-expanderContent {
        max-height: 400px;
        overflow-y: auto;
    }
    </style>
    """,
    unsafe_allow_html=True
)

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    page_bg_img = '''
    <style>
    .stApp {
        background-image: url("data:image/png;base64,%s");
        background-size: cover;
    }
    </style>
    ''' % bin_str
    
    st.markdown(page_bg_img, unsafe_allow_html=True)

# Set the background image
set_png_as_page_bg("background.png")

data_str = {
    'City': ['Tokyo', 'Osaka', 'Kyoto', 'Yokohama', 'Nagoya', 'Sapporo', 'Kobe', 'Fukuoka'],
}

data_int = {
    'Population': [13929286, 2691185, 1464890, 3757630, 2332000, 1970895, 1545873, 1612392],
}

df_sample_str = pd.DataFrame(data_str)
df_sample_int = pd.DataFrame(data_int)
# ================================================ SESSION STATE ============================================================================================================

# Initialize session state
if "page" not in st.session_state:
    st.session_state.page = 0  # 0 for register, 1 for login

if "login" not in st.session_state:
    st.session_state.login = 0  # 0 for not logged in, 1 for logged in

jsonhuhuhaha = "1PlgbJOiKKmOMfn2UH34eNKhuWJRFw7iKveAzuQuKo_w"
url_huhuhaha = f'https://docs.google.com/spreadsheet/ccc?key={jsonhuhuhaha}&output=csv'

if 'credentials_huhuhaha' not in st.session_state:
    df_huhuhaha = pd.read_csv(url_huhuhaha)
    credentials_huhuhaha = df_huhuhaha.iloc[0, 0]
    st.session_state['credentials_huhuhaha'] = credentials_huhuhaha

# Retrieve credentials from session state
credentials_huhuhaha = st.session_state['credentials_huhuhaha']

# Parse the JSON string into a Python dictionary
credentials_dict = json.loads(credentials_huhuhaha)

# Create credentials from the dictionary
creds = service_account.Credentials.from_service_account_info(credentials_dict)

# Add the scopes
scoped_creds = creds.with_scopes(['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive'])

# Authorize the client
client = gspread.authorize(scoped_creds)

# Google Sheets key
sheetskey = "1oJVg7AyQKShgPfCnb_uxfpg1K9nwP0gftv0SILS8BgY"

# ================================================= LOGIN FUNCTIONS =================================================================================================================

def get_login_df():
    """Retrieves login data from Google Sheets."""
    url = f'https://docs.google.com/spreadsheet/ccc?key={sheetskey}&output=csv'
    df_login = pd.read_csv(url)
    return df_login

def send_email(email, message):
    """Sends an email notification."""
    try:
        sender_email = "insightfoxa@gmail.com"  # Replace with your email
        sender_password = "txsr udsn fhpe pfns"  # Replace with your APP PASSWORD

        msg = MIMEText(message)
        msg['Subject'] = 'Password Reminder'
        msg['From'] = sender_email
        msg['To'] = email

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, email, msg.as_string())

        st.success("Reminder email sent, and don't forget to check spam!")
    except Exception as e:
        st.error(f"Failed to send email: {e}")

# ============================================ CLEAN FUNCTIONS =======================================================================================================

def remove_outliers_iqr(df, column):
    """Removes outliers from a DataFrame column using the IQR method."""
    Q1 = df[column].quantile(0.25)
    Q3 = df[column].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    df_filtered = df[(df[column] >= lower_bound) & (df[column] <= upper_bound)]
    return df_filtered

# =============================================== EDA FUNCTIONS ====================================================================================================================

def get_filename_from_title(title):
    """Converts a plot title into a safe file name by replacing spaces and colons."""
    filename = title.replace(" ", "_").replace(":", "") + ".png"
    return filename

# ============================================= Plotting for INT Data ==========================================================================================================================================================================

def get_histogram_figures(df, theme="Blue"):
    """Generates histogram figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#Dcd9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col in enumerate(numerical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sns.histplot(df[col], color=colors[i % len(colors)], kde=True, ax=ax)
            title_str = f"Distribution of {col}"
            ax.set_title(title_str)
            ax.set_xlabel(f"{col} Values")
            ax.set_ylabel("Frequency")
            ax.grid(True, linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating histogram figures: {e}")
        return None

def get_boxplot_figures(df, theme="Blue"):
    """Generates boxplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col in enumerate(numerical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sns.boxplot(x=df[col], color=colors[i % len(colors)], ax=ax)
            title_str = f"Box Plot of {col}"
            ax.set_title(title_str)
            ax.set_xlabel(f"{col} Values")
            ax.grid(True, linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating boxplot figures: {e}")
        return None

def get_scatterplot_figures(df, theme="Blue"):
    """Generates scatterplot figures for unique pairs of numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col1 in enumerate(numerical_cols):
            for j, col2 in enumerate(numerical_cols):
                if i < j:  # Plot only unique pairs
                    fig, ax = plt.subplots(figsize=(8, 6))
                    sns.scatterplot(x=df[col1], y=df[col2], color=colors[(i + j) % len(colors)], ax=ax)
                    title_str = f"Scatter Plot: {col1} vs {col2}"
                    ax.set_title(title_str)
                    ax.set_xlabel(f"{col1} Values")
                    ax.set_ylabel(f"{col2} Values")
                    ax.grid(True, linestyle='--', alpha=0.6)
                    filename = get_filename_from_title(title_str)
                    figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating scatterplot figures: {e}")
        return None

def get_lineplot_figures(df, theme="Blue"):
    """Generates lineplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col in enumerate(numerical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sns.lineplot(x=df.index, y=df[col], color=colors[i % len(colors)], ax=ax)
            title_str = f"Line Plot of {col}"
            ax.set_title(title_str)
            ax.set_xlabel("Index")
            ax.set_ylabel(f"{col} Values")
            ax.grid(True, linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating lineplot figures: {e}")
        return None

def get_areaplot_figures(df, theme="Blue"):
    """Generates areaplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col in enumerate(numerical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.fill_between(df.index, df[col], color=colors[i % len(colors)], alpha=0.6)
            title_str = f"Area Plot of {col}"
            ax.set_title(title_str)
            ax.set_xlabel("Index")
            ax.set_ylabel(f"{col} Values")
            ax.grid(True, linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating areaplot figures: {e}")
        return None

def get_violinplot_figures(df, theme="Blue"):
    """Generates violinplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col in enumerate(numerical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sns.violinplot(y=df[col], color=colors[i % len(colors)], ax=ax)
            title_str = f"Violin Plot of {col}"
            ax.set_title(title_str)
            ax.set_ylabel(f"{col} Values")
            ax.grid(True, linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating violinplot figures: {e}")
        return None

def get_correlation_heatmap_figure(df, theme="Blue"):
    """Generates a correlation heatmap figure for numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#E0F2FE", "#B3E5FC", "#81D4FA", "#4FC3F7", "#29B6F6", "#03A9F4", "#039BE5", "#0288D1", "#01579B"],
            "Green": ["#E8F5E9", "#C8E6C9", "#A5D6A7", "#81C784", "#66BB6A", "#4CAF50", "#43A047", "#388E3C", "#2E7D32"],
            "Red": ["#FFE0E0", "#FFCDD2", "#EF9A9A", "#E57373", "#F44336", "#D32F2F", "#C62828", "#B71C1C", "#8B0000"],
            "Purple": ["#F3E5F5", "#E1BEE7", "#CE93D8", "#BA68C8", "#9C27B0", "#8E24AA", "#7B1FA2", "#6A1B9A", "#4A148C"],
            "Orange": ["#FFF3E0", "#FFE0B2", "#FFCC80", "#FFB366", "#FFA726", "#FF9800", "#FB8C00", "#F57C00", "#E65100"],
            "Gray": ["#F5F5F5", "#EEEEEE", "#E0E0E0", "#BDBDBD", "#9E9E9E", "#757575", "#616161", "#424242", "#212121"],
            "Pastel": ["#FFF9C4", "#FFECB3", "#FFE082", "#FFD54F", "#FFCA28", "#FFC107", "#FFB300", "#FFA000", "#FF6F00"]
        }

        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"

        colors = themes[theme]

        numerical_cols = df.select_dtypes(include=['number']).columns
        corr_matrix = df[numerical_cols].corr()

        fig, ax = plt.subplots(figsize=(10, 8))
        cmap = sns.color_palette(colors, as_cmap=True)
        sns.heatmap(corr_matrix, annot=True, cmap=cmap, linewidths=.5, ax=ax)
        title_str = "Correlation Heatmap"
        ax.set_title(title_str)
        filename = get_filename_from_title(title_str)
        return [(filename, fig)]
    except Exception as e:
        print(f"Error generating correlation heatmap figure: {e}")
        return None

def get_cdf_figures(df, theme="Blue"):
    """Generates CDF figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        numerical_cols = df.select_dtypes(include=['number']).columns
        figures = []
        for i, col in enumerate(numerical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sorted_data = np.sort(df[col])
            yvals = np.arange(len(sorted_data)) / float(len(sorted_data) - 1)
            ax.plot(sorted_data, yvals, color=colors[i % len(colors)])
            title_str = f"CDF of {col}"
            ax.set_title(title_str)
            ax.set_xlabel(f"{col} Values")
            ax.set_ylabel("Cumulative Probability")
            ax.grid(True, linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating CDF figures: {e}")
        return None

# ============================================= Plotting for STR Data ==========================================================================================================================================================================

def get_categorical_barplot_figures(df, theme="Blue"):
    """Generates barplot figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        for i, col in enumerate(categorical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sns.countplot(y=df[col], color=colors[i % len(colors)], ax=ax)
            title_str = f"Bar Plot of {col}"
            ax.set_title(title_str)
            ax.set_xlabel("Count")
            ax.set_ylabel(f"{col} Categories")
            ax.grid(axis='x', linestyle='--', alpha=0.6)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating categorical barplot figures: {e}")
        return None

def get_piechart_figures(df, theme="Blue"):
    """Generates pie chart figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        for i, col in enumerate(categorical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            counts = df[col].value_counts()
            ax.pie(counts, labels=counts.index, autopct='%1.1f%%', colors=colors)
            title_str = f"Pie Chart of {col}"
            ax.set_title(title_str)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating piechart figures: {e}")
        return None

def get_stacked_barplot_figures(df, theme="Blue"):
    """Generates stacked bar plot figures for pairs of categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        if len(categorical_cols) < 2:
            st.write("At least two categorical columns are required for a stacked bar plot.")
            return figures
        for i, col1 in enumerate(categorical_cols):
            for j, col2 in enumerate(categorical_cols):
                if i < j:
                    fig, ax = plt.subplots(figsize=(8, 6))
                    ct = pd.crosstab(df[col1], df[col2])
                    ct.plot(kind='bar', stacked=True, color=colors, ax=ax)
                    title_str = f"Stacked Bar Plot: {col1} vs {col2}"
                    ax.set_title(title_str)
                    ax.set_xlabel(col1)
                    ax.set_ylabel("Count")
                    ax.tick_params(axis='x', rotation=45)
                    plt.tight_layout()
                    filename = get_filename_from_title(title_str)
                    figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating stacked barplot figures: {e}")
        return None

def get_grouped_barplot_figures(df, theme="Blue"):
    """Generates grouped bar plot figures for pairs of categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        if len(categorical_cols) < 2:
            st.write("At least two categorical columns are required for a grouped bar plot.")
            return figures
        for i, col1 in enumerate(categorical_cols):
            for j, col2 in enumerate(categorical_cols):
                if i < j:
                    fig, ax = plt.subplots(figsize=(8, 6))
                    ct = pd.crosstab(df[col1], df[col2])
                    ct.plot(kind='bar', color=colors, ax=ax)
                    title_str = f"Grouped Bar Plot: {col1} vs {col2}"
                    ax.set_title(title_str)
                    ax.set_xlabel(col1)
                    ax.set_ylabel("Count")
                    plt.xticks(rotation=45, ha='right')
                    plt.tight_layout()
                    filename = get_filename_from_title(title_str)
                    figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating grouped barplot figures: {e}")
        return None

def get_wordcloud_figures(df, theme="Blue"):
    """Generates word cloud figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": "Blues",
            "Green": "Greens",
            "Red": "Reds",
            "Purple": "Purples",
            "Orange": "Oranges",
            "Gray": "Greys",
            "Pastel": "Pastel1"
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colormap = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        for i, col in enumerate(categorical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            text = ' '.join(df[col].dropna().astype(str))
            if not text:
                st.write(f"Column '{col}' contains no valid text data.")
                continue
            wordcloud = WordCloud(width=800, height=400, background_color='white', colormap=colormap).generate(text)
            ax.imshow(wordcloud, interpolation='bilinear')
            ax.axis('off')
            title_str = f"Word Cloud of {col}"
            ax.set_title(title_str)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating wordcloud figures: {e}")
        return None

def get_countplot_figures(df, theme="Blue"):
    """Generates count plot figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        for i, col in enumerate(categorical_cols):
            fig, ax = plt.subplots(figsize=(8, 6))
            sns.countplot(data=df, x=col, color=colors[i % len(colors)], ax=ax)
            title_str = f"Count Plot of {col}"
            ax.set_title(title_str)
            ax.set_xlabel(f"{col} Categories")
            ax.set_ylabel("Count")
            ax.tick_params(axis='x', rotation=45)
            plt.tight_layout()
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating countplot figures: {e}")
        return None

def get_treemap_figures(df, theme="Blue"):
    """Generates treemap figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    try:
        themes = {
            "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
            "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
            "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
            "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
            "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
            "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
            "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
        }
        if theme not in themes:
            st.write(f"Theme '{theme}' not found. Using Blue theme.")
            theme = "Blue"
        colors = themes[theme]
        categorical_cols = df.select_dtypes(include=['object']).columns
        figures = []
        for i, col in enumerate(categorical_cols):
            fig, ax = plt.subplots(figsize=(10, 8))
            counts = df[col].value_counts()
            labels = counts.index
            sizes = counts.values
            total = sizes.sum()
            percentages = [f'{size / total * 100:.2f}%' for size in sizes]
            labels_with_percentage = [f'{label}\n({percentage})' for label, percentage in zip(labels, percentages)]
            squarify.plot(sizes=sizes, label=labels_with_percentage, color=colors, alpha=0.8, ax=ax)
            ax.axis('off')
            title_str = f"Treemap of {col}"
            ax.set_title(title_str)
            filename = get_filename_from_title(title_str)
            figures.append((filename, fig))
        return figures
    except Exception as e:
        print(f"Error generating treemap figures: {e}")
        return None

# ============================================= PLOTTING PLOT ONLY ===================================================================================================================================

def plot_histograms(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.histplot(df[col], color=colors[i % len(colors)], kde=True, ax=ax)  # Use ax
        ax.set_title(f"Distribution of {col}")
        ax.set_xlabel(f"{col} Values")
        ax.set_ylabel("Frequency")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_boxplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.boxplot(x=df[col], color=colors[i % len(colors)], ax=ax)  # Use ax
        ax.set_title(f"Box Plot of {col}")
        ax.set_xlabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_scatterplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col1 in enumerate(numerical_cols):
        for j, col2 in enumerate(numerical_cols):
            if i < j:  # Plot only unique pairs
                fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
                sns.scatterplot(x=df[col1], y=df[col2], color=colors[(i + j) % len(colors)], ax=ax)  # Use ax
                ax.set_title(f"Scatter Plot: {col1} vs {col2}")
                ax.set_xlabel(f"{col1} Values")
                ax.set_ylabel(f"{col2} Values")
                ax.grid(True, linestyle='--', alpha=0.6)
                st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_lineplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.lineplot(x=df.index, y=df[col], color=colors[i % len(colors)], ax=ax)  # Use ax
        ax.set_title(f"Line Plot of {col}")
        ax.set_xlabel("Index")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_areaplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        ax.fill_between(df.index, df[col], color=colors[i % len(colors)], alpha=0.6)  # Use ax for plotting
        ax.set_title(f"Area Plot of {col}")
        ax.set_xlabel("Index")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_violinplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.violinplot(y=df[col], color=colors[i % len(colors)], ax=ax)  # Use ax for plotting
        ax.set_title(f"Violin Plot of {col}")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_correlation_heatmap(df, cmap_option="Blue"):
    """
    Plots a correlation heatmap for numerical columns in a Pandas DataFrame using color palettes derived from previous themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot the correlation heatmap from.
        cmap_option (str, optional): Color palette option. Defaults to "Blue".
    """

    cmap_options = {
        "Blue": ["#E0F2FE", "#B3E5FC", "#81D4FA", "#4FC3F7", "#29B6F6", "#03A9F4", "#039BE5", "#0288D1", "#01579B"],
        "Green": ["#E8F5E9", "#C8E6C9", "#A5D6A7", "#81C784", "#66BB6A", "#4CAF50", "#43A047", "#388E3C", "#2E7D32"],
        "Red": ["#FFE0E0", "#FFCDD2", "#EF9A9A", "#E57373", "#F44336", "#D32F2F", "#C62828", "#B71C1C", "#8B0000"],
        "Purple": ["#F3E5F5", "#E1BEE7", "#CE93D8", "#BA68C8", "#9C27B0", "#8E24AA", "#7B1FA2", "#6A1B9A", "#4A148C"],
        "Orange": ["#FFF3E0", "#FFE0B2", "#FFCC80", "#FFB366", "#FFA726", "#FF9800", "#FB8C00", "#F57C00", "#E65100"],
        "Gray": ["#F5F5F5", "#EEEEEE", "#E0E0E0", "#BDBDBD", "#9E9E9E", "#757575", "#616161", "#424242", "#212121"],
        "Pastel": ["#FFF9C4", "#FFECB3", "#FFE082", "#FFD54F", "#FFCA28", "#FFC107", "#FFB300", "#FFA000", "#FF6F00"]
    }

    if cmap_option not in cmap_options:
        print(f"cmap_option '{cmap_option}' not found. Using Blue cmap.")
        cmap_option = "Blue"

    colors = cmap_options[cmap_option]

    numerical_cols = df.select_dtypes(include=['number']).columns
    corr_matrix = df[numerical_cols].corr()

    fig, ax = plt.subplots(figsize=(10, 8))  # Create fig and ax
    cmap = sns.color_palette(colors, as_cmap=True)
    sns.heatmap(corr_matrix, annot=True, cmap=cmap, linewidths=.5, ax=ax)  # Use ax
    ax.set_title("Correlation Heatmap")
    st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_cdf(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax

        # Calculate CDF
        sorted_data = np.sort(df[col])
        yvals = np.arange(len(sorted_data)) / float(len(sorted_data) - 1)

        ax.plot(sorted_data, yvals, color=colors[i % len(colors)])  # Use ax for plotting
        ax.set_title(f"CDF of {col}")
        ax.set_xlabel(f"{col} Values")
        ax.set_ylabel("Cumulative Probability")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass the figure to st.pyplot

# Plotting for STR Data ==========================================================================================================================================================================

def plot_categorical_barplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        sns.countplot(y=df[col], color=colors[i % len(colors)], ax=ax)  # Plot on ax
        ax.set_title(f"Bar Plot of {col}")
        ax.set_xlabel("Count")
        ax.set_ylabel(f"{col} Categories")
        ax.grid(axis='x', linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_piecharts(df, theme="Blue"):
    """
    Plots pie charts for all categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot pie charts from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        counts = df[col].value_counts()
        ax.pie(counts, labels=counts.index, autopct='%1.1f%%', colors=colors)  # Plot on ax
        ax.set_title(f"Pie Chart of {col}")
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_stacked_barplots(df, theme="Blue"):
    """
    Plots stacked bar plots for pairs of categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot stacked bar plots from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    if len(categorical_cols) < 2:
        print("At least two categorical columns are required for a stacked bar plot.")
        return

    for i, col1 in enumerate(categorical_cols):
        for j, col2 in enumerate(categorical_cols):
            if i < j:
                fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
                ct = pd.crosstab(df[col1], df[col2])
                ct.plot(kind='bar', stacked=True, color=colors, ax=ax)  # Plot on ax
                ax.set_title(f"Stacked Bar Plot: {col1} vs {col2}")
                ax.set_xlabel(col1)
                ax.set_ylabel("Count")
                ax.tick_params(axis='x', rotation=45)
                plt.tight_layout()
                st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_grouped_barplots(df, theme="Blue"):
    """
    Plots grouped bar plots (clustered bar plots) for pairs of categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot grouped bar plots from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    if len(categorical_cols) < 2:
        print("At least two categorical columns are required for a grouped bar plot.")
        return

    for i, col1 in enumerate(categorical_cols):
        for j, col2 in enumerate(categorical_cols):
            if i < j:
                fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
                ct = pd.crosstab(df[col1], df[col2])
                ct.plot(kind='bar', color=colors, ax=ax)  # Plot on the ax
                ax.set_title(f"Grouped Bar Plot: {col1} vs {col2}")
                ax.set_xlabel(col1)
                ax.set_ylabel("Count")
                plt.xticks(rotation=45, ha='right')  # Rotate x-axis labels
                plt.tight_layout()
                st.pyplot(fig, clear_figure=True)  # Pass the fig to st.pyplot

def plot_wordclouds(df, theme="Blue"):
    """
    Plots word clouds for all categorical (string) columns in a Pandas DataFrame with color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot word clouds from.
        theme (str, optional): The color theme for the word clouds. Defaults to "Blue".
    """

    themes = {
        "Blue": "Blues",
        "Green": "Greens",
        "Red": "Reds",
        "Purple": "Purples",
        "Orange": "Oranges",
        "Gray": "Greys",
        "Pastel": "Pastel1"
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colormap = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        text = ' '.join(df[col].dropna().astype(str))  # Combine all strings into one text
        if not text:
            print(f"Column '{col}' contains no valid text data.")
            continue

        wordcloud = WordCloud(width=800, height=400, background_color='white', colormap=colormap).generate(text)
        ax.imshow(wordcloud, interpolation='bilinear')  # Plot on ax
        ax.axis('off')
        ax.set_title(f"Word Cloud of {col}")
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_countplots(df, theme="Blue"):
    """
    Plots count plots for all categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot count plots from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        sns.countplot(data=df, x=col, color=colors[i % len(colors)], ax=ax)  # Use ax
        ax.set_title(f"Count Plot of {col}")
        ax.set_xlabel(f"{col} Categories")
        ax.set_ylabel("Count")
        ax.tick_params(axis='x', rotation=45)
        plt.tight_layout()
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_treemaps(df, theme="Blue"):
    """
    Plots treemaps for categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes, including percentage values on each block.

    Args:
        df (pd.DataFrame): The DataFrame to plot treemaps from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(10, 8))  # Create fig, ax
        counts = df[col].value_counts()
        labels = counts.index
        sizes = counts.values

        # Calculate percentages
        total = sizes.sum()
        percentages = [f'{size / total * 100:.2f}%' for size in sizes]

        # Combine labels with percentages
        labels_with_percentage = [f'{label}\n({percentage})' for label, percentage in zip(labels, percentages)]

        squarify.plot(sizes=sizes, label=labels_with_percentage, color=colors, alpha=0.8, ax=ax)  # Use ax for plotting
        ax.axis('off')
        ax.set_title(f"Treemap of {col}")
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

# ========================================== REPORT VARIATIONS LIST =============================================================================================================

# English 

grouping_obs_intro_en = [
    "Analyzing the data by monthly segments across multiple years reveals the period with the most frequent observations.",
    "When data is segmented into monthly intervals over several years, the peak observation frequency is pinpointed.",
    "Segmenting the data into monthly blocks across years allows us to identify the interval with the highest number of observations.",
    "Through a monthly segmentation of data over years, the interval showcasing the greatest observation frequency becomes clear.",
    "We can determine the interval with the highest observation rate by organizing the data into monthly segments spanning multiple years.",
    "The data, when broken down into monthly segments across years, highlights the period with the maximum observation frequency.",
    "By dividing the data into month-based segments over a multi-year period, the interval with the highest observation count is determined."
]

seasonal_interpretation_en = [
    "The seasonal patterns we see are likely caused by events that happen regularly, such as weather changes, holiday seasons, or other repeating influences.",
    "It's probable that the seasonal fluctuations are driven by recurring events, encompassing climatic patterns, holiday rhythms, and other periodic elements.",
    "Observed seasonal variations are most likely attributable to events that repeat, whether they are due to climate, holidays, or other cyclical factors.",
    "The seasonal variation observed is probably due to events that occur repeatedly, including weather patterns, holiday schedules, and other periodic variables.",
    "We can infer that the seasonal changes are a result of repeating events, which could be related to weather, holidays, or other regular cycles.",
    "Recurring events, such as climatic conditions, holiday schedules, or other periodic factors, are likely responsible for the seasonal patterns observed.",
    "The seasonal variations are likely the result of repeated events, be they due to climate, holiday periods, or other regularly occurring influences."
]

advanced_ts_intro_en = [
    "A thorough examination of the time series reveals specific insights.",
    "By conducting a detailed analysis of the time series, significant findings emerge.",
    "Upon a meticulous investigation of the time series, it becomes evident that certain patterns exist.",
    "In-depth exploration of the time series uncovers key characteristics.",
    "A careful and extensive analysis of the time series data provides crucial discoveries.",
    "Through a rigorous analytical process, the time series data reveals important trends.",
    "Detailed scrutiny of the time series demonstrates that specific features are present."
]

descriptive_stat_intro_en = [
    "A quick look at the statistical summary provides key information.",
    "The statistical summary offers a brief overview of the data's characteristics.",
    "We can gain a basic understanding of the data through a concise statistical summary.",
    "A short summary of the statistical data gives us a general idea of its properties.",
    "The data's statistical highlights are captured in a brief summary.",
    "A snapshot of the statistics reveals essential details about the data.",
    "A preliminary overview of the statistics presents the data's fundamental attributes."
]

univariate_analysis_intro_en = [
    "Detailed single-variable metrics offer insights into the data's characteristics.",
    "Comprehensive analysis of individual variables shows specific trends.",
    "Examining each variable separately provides detailed statistical measures.",
    "In-depth metrics for individual variables reveal key aspects of the data.",
    "A thorough examination of single variables yields detailed insights.",
    "Analyzing each variable in isolation provides detailed statistical information.",
    "By focusing on single variables, we can obtain detailed measurement data."
]

bivariate_analysis_intro_en = [
    "Comparing two variables suggests specific relationships.",
    "The analysis of two variables together indicates potential associations.",
    "Evaluating the relationship between two variables reveals possible connections.",
    "An examination of two variables side by side provides evidence of potential correlations.",
    "Looking at two variables simultaneously suggests certain patterns.",
    "Assessing the interaction between two variables indicates potential links.",
    "The joint analysis of two variables highlights potential relationships."
]

rolling_stats_intro_en = [
    "Calculations of moving averages reveal specific trends in the data.",
    "The data's smoothed trends are evident from the computed rolling averages.",
    "By calculating rolling averages, we can observe underlying patterns in the data.",
    "The analysis of moving averages highlights certain shifts and changes over time.",
    "Computed rolling statistics demonstrate the presence of specific trends.",
    "Utilizing rolling averages allows for the identification of trends within the dataset.",
    "The trends revealed by the rolling average calculations are significant."
]

time_range_context_en = [
    "The period from {start} to {end} is crucial for understanding both immediate changes and broader patterns.",
    "Analyzing the data between {start} and {end} provides key insights into both temporary and long-term variations.",
    "Examining the data spanning {start} to {end} allows us to observe both short-term variability and long-term trends.",
    "The timeframe from {start} to {end} provides essential information about the data's dynamic behavior.",
    "Within the interval of {start} to {end}, we can discern significant short-term fluctuations and long-term trends.",
    "Observations within the {start} to {end} period are critical for understanding the data's immediate and sustained changes.",
    "Understanding the data's behavior over the period from {start} to {end} is key to identifying both immediate and gradual changes."
]

overall_insight_intro_en = [
    "In summary, the analysis clearly demonstrates specific findings.",
    "To conclude, the analysis provides definitive evidence of certain conclusions.",
    "Essentially, the analysis proves specific points conclusively.",
    "Briefly stated, the analysis provides clear and final results.",
    "In essence, the data analysis provides conclusive insights.",
    "To put it simply, the analysis leads to clear and decisive conclusions.",
    "In short, the analysis demonstrates definitive outcomes."
]

variables_detail_intro_en = [
    "A thorough assessment of this variable confirms specific characteristics.",
    "Detailed study of this variable provides confirmation of certain attributes.",
    "By conducting an in-depth evaluation, we can confirm the variable's specific properties.",
    "The detailed evaluation of this variable validates specific findings.",
    "An extensive analysis of this variable confirms particular aspects.",
    "Through a meticulous review, the variable's characteristics are confirmed.",
    "The detailed analysis supports the specific qualities of this variable."
]

data_structure_context_en = [
    "Using the diverse structure of the data enables us to apply specific analytical methods efficiently.",
    "Leveraging the dataset's varied structure allows for the effective use of targeted analytical strategies.",
    "The dataset's structural diversity facilitates the application of appropriate analytical techniques.",
    "By taking advantage of the data's structural variety, we can apply specialized analytical approaches.",
    "The varied structure of the dataset allows for the effective deployment of specific analytical methods.",
    "Utilizing the dataset's structural complexity enables the application of focused analytical techniques.",
    "The diverse data structure supports the use of tailored analytical strategies for different data types."
]

observation_counts_summary_en = [
    "Analyzing the data by time segments shows that observation numbers vary significantly across different periods.",
    "Segmenting the data based on time reveals substantial differences in observation counts between periods.",
    "The time-based segmentation of data highlights the variability in observation counts over time.",
    "By segmenting the data chronologically, we observe considerable fluctuations in observation counts.",
    "The division of data by time intervals indicates that observation counts differ greatly between periods.",
    "A temporal segmentation of the data reveals significant variations in the number of observations across different time frames.",
    "When data is segmented by time, it becomes clear that observation counts are not consistent across all periods."
]

rolling_trend_summary_en = [
    "The ongoing trends are confirmed by the rolling analysis throughout the observed timeframe.",
    "Analysis using rolling statistics validates the consistent trends seen during the observation period.",
    "Rolling computations demonstrate that the trends remain consistent over the entire period analyzed.",
    "The data's trends are shown to be persistent through the rolling analysis conducted on the observed period.",
    "By applying rolling analysis, the continued presence of trends is confirmed across the observed data.",
    "Consistent trends over the observed period are supported by the findings of the rolling analysis.",
    "The persistence of trends within the data is verified by the rolling analysis across the entire observation period."
]

time_series_freq_explanation_en = [
    "Based on frequency analysis, it can be inferred that the data is collected in a consistent, repeating pattern.",
    "Frequency inference reveals that the data is collected according to a predictable, periodic cycle.",
    "The data's regular collection pattern is suggested by the frequency analysis conducted.",
    "Through frequency inference, we can deduce that the data follows a routine, cyclical collection process.",
    "Evidence suggests, based on frequency analysis, that the data is gathered at regular, periodic intervals.",
    "Analysis of the data's frequency indicates a consistent and recurring collection cycle.",
    "It is clear from the frequency inference that the data adheres to a regular, periodic collection schedule."
]

final_overall_summary_en = [
    "Ultimately, the combined findings offer a strong basis for future analytical strategies and model development.",
    "In conclusion, the compiled results provide a firm foundation for the direction of subsequent data analysis and model creation.",
    "To summarize, the integrated findings establish a robust platform for the progression of data analysis and model building.",
    "As a final point, the consolidated findings create a solid groundwork for the future of data analysis and model development.",
    "In closing, the collective findings give a stable base for the future path of data analysis and the creation of models.",
    "To conclude, the synthesized findings are a dependable starting point for the evolution of data analysis and model design.",
    "Finally, the combined results present a strong base for future data exploration and model construction."
]

mode_std_variance_summary_en = [
    "Given the mode of {mode_val}, the data exhibits a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, indicating significant variability within the dataset.",
    "With {mode_val} as the most frequent value, the data shows a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, implying considerable dispersion.",
    "The data, having a mode of {mode_val}, demonstrates a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, which suggests a notable degree of variability.",
    "A mode of {mode_val} is observed, and the dataset's standard deviation is {std_val:.4f} and variance is {var_val:.4f}, highlighting substantial variability.",
    "Having a mode of {mode_val}, the data reveals a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, indicating a considerable amount of variation.",
    "The data's most common value is {mode_val}, and it shows a standard deviation of {std_val:.4f} and variance of {var_val:.4f}, pointing to significant variability.",
    "With {mode_val} as the mode, the data has a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, suggesting a high level of variability."
]

range_iqr_summary_en = [
    "The data analysis reveals an overall range of {range_val:.4f} and an interquartile range (IQR) of {iqr_val:.4f}, both of which collectively illustrate the spread of the data.",
    "Analysis of the data shows a total range of {range_val:.4f} and an IQR of {iqr_val:.4f}, emphasizing the dispersion of the dataset.",
    "Data analysis indicates that the dataset has a range of {range_val:.4f} and an IQR of {iqr_val:.4f}, which together demonstrate the data's spread.",
    "An overall range of {range_val:.4f} and an IQR of {iqr_val:.4f} are observed in the data analysis, highlighting the dataset's distribution.",
    "The data's spread is evidenced by its range of {range_val:.4f} and its IQR of {iqr_val:.4f}, as shown by the analysis.",
    "The dataset's dispersion is characterized by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}, according to the data analysis.",
    "Analysis of the data demonstrates a spread as indicated by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}."
]

overall_dispersion_summary_en = [
    "In general, the considerable dispersion in the data is highlighted by the high values of standard deviation ({std_val:.4f}) and variance ({var_val:.4f}), along with a range of {range_val:.4f}.",
    "The data's notable spread is evident from the high standard deviation ({std_val:.4f}) and variance ({var_val:.4f}), as well as the range of {range_val:.4f}.",
    "Overall, a substantial degree of dispersion is demonstrated by the large standard deviation ({std_val:.4f}) and variance ({var_val:.4f}), in conjunction with a range of {range_val:.4f}.",
    "The dataset's dispersion is clearly shown by the elevated standard deviation ({std_val:.4f}) and variance ({var_val:.4f}), along with the range of {range_val:.4f}.",
    "Remarkable dispersion is indicated by the high standard deviation ({std_val:.4f}) and variance ({var_val:.4f}), as well as the range of {range_val:.4f} observed in the data.",
    "The data's significant dispersion is reflected in the high standard deviation ({std_val:.4f}), variance ({var_val:.4f}), and range of {range_val:.4f}.",
    "High values for standard deviation ({std_val:.4f}) and variance ({var_val:.4f}), coupled with a range of {range_val:.4f}, collectively signify considerable data dispersion."
]

distribution_metrics_summary_en = [
    "The descriptive analysis highlights that the data has a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, alongside a range of {range_val:.4f} and an interquartile range (IQR) of {iqr_val:.4f}.",
    "From the descriptive analysis, we observe a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, with the data's spread being represented by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}.",
    "A mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val} are revealed by the descriptive analysis, with the data's dispersion quantified by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}.",
    "The descriptive analysis indicates that the dataset's central tendencies include a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, while its variability is captured by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}.",
    "Analysis of the descriptive statistics shows a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, with the data's spread defined by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}.",
    "Descriptive metrics show the data's mean to be {mean_val:.4f}, its median {median_val:.4f}, and its mode {mode_val}, and its spread is further defined by a range of {range_val:.4f} and an IQR of {iqr_val:.4f}.",
    "The data's distribution characteristics, as revealed by descriptive analysis, include a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a range of {range_val:.4f}, and an IQR of {iqr_val:.4f}."
]

central_tendency_dispersion_en = [
    "The overall distribution is depicted by central measures such as a mode of {mode_val} and a median of {median_val:.4f}, along with dispersion measures like a standard deviation of {std_val:.4f} and a range of {range_val:.4f}.",
    "Central measures, including a mode of {mode_val} and a median of {median_val:.4f}, combined with measures of spread like a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, illustrate the data's overall distribution.",
    "The data's distribution is characterized by a mode of {mode_val} and a median of {median_val:.4f}, which reflect central tendencies, and a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, which indicate dispersion.",
    "A mode of {mode_val} and a median of {median_val:.4f} provide insights into the data's central tendencies, while a standard deviation of {std_val:.4f} and a range of {range_val:.4f} describe its dispersion, together showing the overall distribution.",
    "The overall distribution is represented by a mode of {mode_val} and a median of {median_val:.4f}, highlighting central tendencies, and a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, showing data spread.",
    "Central tendencies, as shown by a mode of {mode_val} and a median of {median_val:.4f}, and dispersion, represented by a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, together define the data's distribution.",
    "The data's overall distribution is captured by its central measures, a mode of {mode_val} and a median of {median_val:.4f}, and its dispersion measures, a standard deviation of {std_val:.4f} and a range of {range_val:.4f}."
]

descriptive_stats_template_en = [
    "The mean and median, measured at {mean_val:.4f} and {median_val:.4f} respectively, indicate a {distribution} distribution in the data.",
    "A mean of {mean_val:.4f} and a median of {median_val:.4f} are observed, suggesting that the data follows a {distribution} distribution pattern.",
    "Descriptive metrics reveal a mean of {mean_val:.4f} and a median of {median_val:.4f}, which together point to a {distribution} distribution.",
    "Given the mean of {mean_val:.4f} and the median of {median_val:.4f}, the data's distribution is classified as {distribution}.",
    "The data exhibits a mean of {mean_val:.4f} and a median of {median_val:.4f}, indicating that its distribution can be characterized as {distribution}.",
    "With a mean of {mean_val:.4f} and a median of {median_val:.4f}, the data's distribution is suggestive of a {distribution} pattern.",
    "The descriptive statistical measures, a mean of {mean_val:.4f} and a median of {median_val:.4f}, are indicative of a {distribution} distribution."
]

variability_insight_en = [
    "The data's dispersion is significant, with a standard deviation of {std_val:.4f} and a variance of {var_val:.4f} indicating a high level of spread.",
    "A notable degree of variability is observed in the data, as evidenced by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "The data demonstrates considerable spread, highlighted by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "High data spread is indicated by the standard deviation of {std_val:.4f} and the variance of {var_val:.4f}.",
    "The data's high spread is quantified by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "Significant variability in the data is shown by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "The data's dispersion is substantial, as shown by its standard deviation of {std_val:.4f} and variance of {var_val:.4f}."
]

statistical_measure_overview_en = [
    "An overview of the statistical measures shows a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, suggesting a broad distribution of values.",
    "The statistical assessment reveals a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, indicating that the data is widely dispersed.",
    "A broad distribution of values is suggested by the statistical overview, which includes a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "The data's values are widely distributed, as shown by the statistical assessment, which reports a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "Overall, the statistical measures, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, point to a wide spread in the data.",
    "The data's distribution is wide, as indicated by the statistical overview, which provides a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "A wide distribution of values is evident from the statistical measure overview, which includes a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}."
]

central_tendency_dispersion_en = [
    "The data's overall distribution is visualized through central measures, such as a mode of {mode_val} and a median of {median_val:.4f}, along with measures of dispersion, including a standard deviation of {std_val:.4f} and a range of {range_val:.4f}.",
    "Central tendencies, indicated by a mode of {mode_val} and a median of {median_val:.4f}, and data spread, shown by a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, together provide a picture of the overall distribution.",
    "The distribution of the data is characterized by central measures, a mode of {mode_val} and a median of {median_val:.4f}, and measures of variability, a standard deviation of {std_val:.4f} and a range of {range_val:.4f}.",
    "A comprehensive view of the data's distribution is offered by central measures, a mode of {mode_val} and a median of {median_val:.4f}, and dispersion measures, a standard deviation of {std_val:.4f} and a range of {range_val:.4f}.",
    "The data's distribution is represented by its central tendencies, a mode of {mode_val} and a median of {median_val:.4f}, and its spread, a standard deviation of {std_val:.4f} and a range of {range_val:.4f}.",
    "Central measures, including a mode of {mode_val} and a median of {median_val:.4f}, and dispersion measures, such as a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, jointly illustrate the data's overall distribution.",
    "The data's distribution is thoroughly described by its central measures, a mode of {mode_val} and a median of {median_val:.4f}, and its variability, a standard deviation of {std_val:.4f} and a range of {range_val:.4f}."
]

variability_insight_en = [
    "The data's dispersion is significant, with a standard deviation of {std_val:.4f} and a variance of {var_val:.4f} indicating a high level of spread.",
    "A notable degree of variability is observed in the data, as evidenced by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "The data demonstrates considerable spread, highlighted by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "High data spread is indicated by the standard deviation of {std_val:.4f} and the variance of {var_val:.4f}.",
    "The data's high spread is quantified by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "Significant variability in the data is shown by a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}.",
    "The data's dispersion is substantial, as shown by its standard deviation of {std_val:.4f} and variance of {var_val:.4f}."
]

statistical_measure_overview_en = [
    "A review of the essential statistical metrics reveals a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, suggesting a broad dispersion of values.",
    "The key statistical measures, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, indicate a wide spread in the data.",
    "Examining the core statistical measures, a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, highlights the data's wide distribution.",
    "The data's wide spread is evident from the review of key statistics, which includes a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "A wide distribution of values is indicated by the statistical review, which presents a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "The review of critical statistical measures, a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, points to a wide spread in the data.",
    "The data's dispersion is clearly shown in the statistical overview, which reports a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}."
]

spread_and_skew_summary_en = [
    "The distribution's variability and asymmetry are highlighted by a standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f}.",
    "A standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f} together illustrate the data's dispersion and lack of symmetry.",
    "The data's spread and its deviation from symmetry are captured by a standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f}.",
    "Measured variability and asymmetry in the distribution are shown by a standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f}.",
    "The distribution's characteristics, in terms of spread and skew, are defined by a standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f}.",
    "A standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f} provide insights into the data's dispersion and asymmetry.",
    "The variability and asymmetry of the data's distribution are clearly indicated by a standard deviation of {std_val:.4f} and a skewness value of {skew_val:.4f}."
]

outlier_context_summary_en = [
    "The presence of {outlier_count} identified outliers suggests the necessity for further investigation to ensure data integrity.",
    "An outlier count of {outlier_count} indicates that additional analysis is required to assess the data's reliability.",
    "With {outlier_count} outliers detected, it is recommended to conduct further investigation to verify the data's accuracy.",
    "The detection of {outlier_count} outliers implies that further scrutiny is needed to confirm the data's trustworthiness.",
    "The identified {outlier_count} outliers necessitate additional investigation to evaluate the data's integrity.",
    "An outlier count of {outlier_count} calls for further assessment to determine the data's reliability.",
    "The presence of {outlier_count} outliers suggests that a more thorough investigation is needed to ensure the data's validity."
]

variance_range_iqr_detail_en = [
    "The data's overall spread is effectively captured by a variance of {var_val:.4f}, a range of {range_val:.4f}, and an interquartile range (IQR) of {iqr_val:.4f}, all of which indicate a wide dispersion.",
    "Broad dispersion is indicated by a variance of {var_val:.4f}, a range of {range_val:.4f}, and an IQR of {iqr_val:.4f}, all of which represent the data's overall spread.",
    "A variance of {var_val:.4f}, a range of {range_val:.4f}, and an IQR of {iqr_val:.4f} collectively illustrate the data's wide dispersion and overall spread.",
    "The overall spread of the data, as shown by a variance of {var_val:.4f}, a range of {range_val:.4f}, and an IQR of {iqr_val:.4f}, points to a broad dispersion.",
    "Data dispersion is comprehensively represented by a variance of {var_val:.4f}, a range of {range_val:.4f}, and an IQR of {iqr_val:.4f}, all of which highlight the data's overall spread.",
    "The data's broad dispersion is indicated by a variance of {var_val:.4f}, a range of {range_val:.4f}, and an IQR of {iqr_val:.4f}, which together capture its overall spread.",
    "Overall data spread is thoroughly represented by a variance of {var_val:.4f}, a range of {range_val:.4f}, and an interquartile range of {iqr_val:.4f}, indicating broad dispersion."
]

mode_and_dispersion_detail_en = [
    "The most frequently occurring value in the data is {mode_val}, and this, combined with a standard deviation of {std_val:.4f}, points to a significant spread within the data.",
    "With {mode_val} as the mode, the data exhibits a considerable spread, as evidenced by a standard deviation of {std_val:.4f}.",
    "The data's mode, {mode_val}, when considered alongside a standard deviation of {std_val:.4f}, reveals a notable degree of dispersion.",
    "A standard deviation of {std_val:.4f} indicates significant data spread, especially when the mode is {mode_val}.",
    "The data's spread is substantial, as shown by a standard deviation of {std_val:.4f}, given that the most common value is {mode_val}.",
    "The mode, {mode_val}, in conjunction with a standard deviation of {std_val:.4f}, highlights the considerable spread of the data.",
    "Significant dispersion within the data is suggested by a standard deviation of {std_val:.4f}, particularly when the mode is {mode_val}."
]

statistical_overview_paraphrase_en = [
    "The overall statistical assessment reveals a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, indicating that the data's values are widely distributed.",
    "A wide distribution of values is suggested by the statistical overview, which reports a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "The data's dispersion is evident from the statistical assessment, showing a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "An examination of the overall statistics, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, points to a broad spread in the data.",
    "The statistical overview highlights a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, which together indicate a wide distribution of values.",
    "A mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, as shown in the statistical assessment, suggest that the data's values are spread out.",
    "The data's wide distribution is confirmed by the statistical overview, which includes a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}."
]

univariate_overview_metrics_en = [
    "The univariate analysis presents key statistical features of the variable, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}.",
    "A summary of the variable's statistical characteristics, such as a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, is provided by the univariate analysis.",
    "The univariate analysis yields fundamental statistical measures, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, effectively summarizing the variable.",
    "Key statistical metrics, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, are obtained through univariate analysis, offering insights into the variable's behavior.",
    "Univariate analysis results in a set of core statistical values—a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}—that succinctly describe the variable.",
    "The variable's essential statistical properties are encapsulated in the mean of {mean_val:.4f}, the median of {median_val:.4f}, and the mode of {mode_val}, as determined by univariate analysis.",
    "A comprehensive summary of the variable's statistics, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, is achieved through univariate analysis."
]

detailed_statistical_insight_en = [
    "The variable's statistical profile, including a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}, along with a standard deviation of {std_val:.4f} and a range of {range_val:.4f}, is thoroughly confirmed by the detailed statistical analysis.",
    "Detailed statistical analysis provides confirmation of the variable's central tendencies and dispersion, as represented by a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a standard deviation of {std_val:.4f}, and a range of {range_val:.4f}.",
    "The variable's distribution characteristics, encompassing a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a standard deviation of {std_val:.4f}, and a range of {range_val:.4f}, are validated by the detailed statistical analysis.",
    "A comprehensive picture of the variable's statistical properties, including its central tendencies and dispersion, is offered by the detailed analysis, which confirms a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a standard deviation of {std_val:.4f}, and a range of {range_val:.4f}.",
    "The detailed statistical analysis substantiates the variable's statistical features, which include a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a standard deviation of {std_val:.4f}, and a range of {range_val:.4f}.",
    "The variable's statistical makeup, as defined by a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a standard deviation of {std_val:.4f}, and a range of {range_val:.4f}, is confirmed through detailed statistical analysis.",
    "Through detailed statistical analysis, the variable is shown to have a mean of {mean_val:.4f}, a median of {median_val:.4f}, a mode of {mode_val}, a standard deviation of {std_val:.4f}, and a range of {range_val:.4f}, all of which collectively depict its distribution."
]

time_series_analysis_intro_en = [
    "The time series component spans from {start_date} to {end_date}. This extended period allows us to capture long-term trends, subtle cycles, and abrupt shifts that are critical for forecasting.",
    "Covering the period from {start_date} to {end_date}, the dataset provides a comprehensive temporal context that enables the detection of persistent trends, recurring seasonal patterns, and occasional anomalies.",
    "Spanning from {start_date} to {end_date}, the analysis benefits from a lengthy observation window, offering insights into both enduring trends and transient irregularities that can guide strategic decisions.",
    "From {start_date} to {end_date}, the dataset offers a robust temporal framework, helping to uncover gradual trends, periodic fluctuations, and sporadic disruptions throughout the observed period.",
    "The observation period, stretching from {start_date} to {end_date}, provides a solid basis for time series analysis, capturing both long-term evolution and short-term variability in the data.",
    "Analyzing data from {start_date} to {end_date} reveals a rich temporal context, enabling the identification of sustained trends, seasonal cycles, and momentary deviations that inform robust forecasting."
]

# 2. Seasonal Pattern Analysis Insights (for each numeric variable resampled by user-defined frequency)
seasonal_pattern_insight_en = [
    "For '{col}', resampling at the user-defined frequency reveals that the highest average is {max_val:.4f} during {max_period} (which is {max_status} the overall average),\nand the lowest average is {min_val:.4f} during {min_period} (which is {min_status} the overall average). This clearly indicates a recurring seasonal cycle.",
    "Examining '{col}', the resampled data shows a peak average of {max_val:.4f} in the period {max_period} ({max_status} the overall mean),\nand a trough average of {min_val:.4f} in the period {min_period} ({min_status} the overall mean), suggesting pronounced seasonal fluctuations.",
    "The seasonal analysis for '{col}' indicates that the period {max_period} records the highest average at {max_val:.4f} ({max_status} the overall average),\nwhile the period {min_period} has the lowest average at {min_val:.4f} ({min_status} the overall average), highlighting clear seasonal behavior.",
    "For '{col}', the data resampled by {timefreq} reveals distinct seasonal trends: the maximum average of {max_val:.4f} occurs in {max_period} ({max_status} average),\nand the minimum average of {min_val:.4f} appears in {min_period} ({min_status} average).",
    "Analysis of '{col}' shows that when aggregated by the selected frequency, the highest average value is {max_val:.4f} in {max_period} ({max_status} the mean),\nand the lowest average value is {min_val:.4f} in {min_period} ({min_status} the mean), underscoring significant seasonal variation.",
    "For '{col}', seasonal grouping by {timefreq} indicates a peak average of {max_val:.4f} during {max_period} ({max_status} the overall mean) and a trough of {min_val:.4f} during {min_period} ({min_status} the overall mean), clearly depicting seasonal trends."
]

# 3. Rolling Statistics and Trend Analysis Insights (for each numeric variable's rolling metrics)
rolling_trend_insight_en = [
    "For '{col}', a rolling window of 5 periods reveals an average rolling mean of {roll_mean:.4f} and an average rolling standard deviation of {roll_std:.4f}.\nThis indicates a relatively stable trend with periodic fluctuations that may highlight transitional phases in the data.",
    "The rolling analysis for '{col}' (using a 5-period window) shows a moving average of {roll_mean:.4f} alongside a rolling standard deviation of {roll_std:.4f}.\nThese metrics suggest that while the overall trend remains consistent, there are intermittent variations worth further exploration.",
    "Evaluating '{col}' over a rolling window of 5 periods yields an average rolling mean of {roll_mean:.4f} and a standard deviation of {roll_std:.4f}.\nSuch patterns indicate that the variable maintains a stable trend interspersed with occasional volatility, pointing to potential shifts in behavior.",
    "For '{col}', the computed rolling statistics using a window of 5 periods result in an average moving mean of {roll_mean:.4f} and an average standard deviation of {roll_std:.4f}.\nThis analysis reflects consistent trends with periodic deviations that could signify seasonal or cyclical events.",
    "A 5-period rolling analysis of '{col}' indicates an average rolling mean of {roll_mean:.4f} and an average rolling standard deviation of {roll_std:.4f}.\nThese findings highlight a stable trend punctuated by moderate fluctuations, providing insights into the variable’s underlying dynamics.",
    "The rolling metrics for '{col}', calculated over 5 periods, show an average rolling mean of {roll_mean:.4f} and a rolling standard deviation of {roll_std:.4f}.\nThis suggests that while the overall trend is stable, there are consistent, periodic variations that may point to underlying cyclical patterns."
]


# Indonesian 

grouping_obs_intro_id = [
    "Dengan membagi data menjadi segmen bulanan selama beberapa tahun, kita dapat menentukan periode dengan frekuensi pengamatan tertinggi.",
    "Pengelompokan data berdasarkan bulan selama bertahun-tahun memungkinkan kita untuk mengidentifikasi interval dengan jumlah pengamatan terbanyak.",
    "Melalui segmentasi data bulanan lintas tahun, periode dengan frekuensi pengamatan maksimum menjadi jelas.",
    "Data yang dikategorikan dalam segmen bulanan selama beberapa tahun membantu kita menemukan interval dengan tingkat pengamatan tertinggi.",
    "Kita dapat mengetahui interval dengan frekuensi pengamatan tertinggi dengan mengatur data ke dalam segmen bulanan yang mencakup beberapa tahun.",
    "Pembagian data menjadi segmen berbasis bulan selama periode multi-tahun menentukan interval dengan jumlah pengamatan terbanyak.",
    "Analisis data per segmen bulanan dari tahun ke tahun memperlihatkan periode dengan jumlah pengamatan paling sering."
]

seasonal_interpretation_id = [
    "Pola musiman yang kita lihat kemungkinan besar disebabkan oleh kejadian yang berulang, seperti perubahan cuaca, musim liburan, atau faktor periodik lainnya.",
    "Kemungkinan besar fluktuasi musiman dipicu oleh peristiwa berulang, termasuk pola iklim, ritme liburan, dan elemen periodik lainnya.",
    "Variasi musiman yang diamati kemungkinan besar disebabkan oleh peristiwa yang terjadi berulang kali, baik itu karena iklim, liburan, atau faktor siklus lainnya.",
    "Peristiwa berulang, seperti kondisi iklim, jadwal liburan, atau faktor periodik lainnya, kemungkinan besar bertanggung jawab atas pola musiman yang diamati.",
    "Kita dapat menyimpulkan bahwa perubahan musiman adalah hasil dari peristiwa berulang, yang bisa terkait dengan cuaca, liburan, atau siklus reguler lainnya.",
    "Variasi musiman kemungkinan besar merupakan hasil dari kejadian berulang, baik itu karena iklim, periode liburan, atau pengaruh yang terjadi secara teratur.",
    "Peristiwa yang terjadi secara periodik seperti cuaca, liburan, atau kejadian lain yang berulang, merupakan penyebab dari pola musiman yang terlihat."
]

advanced_ts_intro_id = [
    "Pemeriksaan mendalam terhadap deret waktu mengungkapkan wawasan spesifik.",
    "Melalui analisis rinci deret waktu, temuan signifikan muncul.",
    "Setelah penyelidikan cermat terhadap deret waktu, menjadi jelas bahwa pola tertentu ada.",
    "Eksplorasi mendalam deret waktu mengungkap karakteristik utama.",
    "Analisis yang cermat dan ekstensif terhadap data deret waktu memberikan penemuan penting.",
    "Melalui proses analitis yang ketat, data deret waktu mengungkapkan tren penting.",
    "Penelitian mendalam terhadap deret waktu menunjukkan bahwa fitur spesifik hadir."
]

descriptive_stat_intro_id = [
    "Tinjauan singkat ringkasan statistik memberikan informasi penting.",
    "Ringkasan statistik menawarkan gambaran singkat tentang karakteristik data.",
    "Kita dapat memperoleh pemahaman dasar tentang data melalui ringkasan statistik yang ringkas.",
    "Ringkasan singkat data statistik memberi kita gambaran umum tentang propertinya.",
    "Sorotan statistik data ditangkap dalam ringkasan singkat.",
    "Cuplikan statistik mengungkapkan detail penting tentang data.",
    "Tinjauan awal statistik menyajikan atribut dasar data."
]

univariate_analysis_intro_id = [
    "Metrik variabel tunggal terperinci menawarkan wawasan tentang karakteristik data.",
    "Analisis komprehensif variabel individu menunjukkan tren spesifik.",
    "Memeriksa setiap variabel secara terpisah memberikan ukuran statistik terperinci.",
    "Metrik mendalam untuk variabel individu mengungkapkan aspek kunci data.",
    "Pemeriksaan menyeluruh terhadap variabel tunggal menghasilkan wawasan terperinci.",
    "Menganalisis setiap variabel secara terpisah memberikan informasi statistik terperinci.",
    "Dengan berfokus pada variabel tunggal, kita dapat memperoleh data pengukuran terperinci."
]

bivariate_analysis_intro_id = [
    "Membandingkan dua variabel menunjukkan hubungan spesifik.",
    "Analisis dua variabel bersama-sama menunjukkan potensi hubungan.",
    "Mengevaluasi hubungan antara dua variabel mengungkapkan kemungkinan koneksi.",
    "Pemeriksaan dua variabel berdampingan memberikan bukti potensi korelasi.",
    "Melihat dua variabel secara bersamaan menunjukkan pola tertentu.",
    "Menilai interaksi antara dua variabel menunjukkan potensi hubungan.",
    "Analisis gabungan dari dua variabel menyoroti potensi hubungan."
]

rolling_stats_intro_id = [
    "Perhitungan rata-rata bergerak telah dilakukan, menunjukkan adanya tren tertentu dalam data.",
    "Tren data yang dihaluskan terlihat dari perhitungan rata-rata bergerak.",
    "Dengan menghitung rata-rata bergerak, kita dapat mengamati pola yang mendasari data.",
    "Analisis rata-rata bergerak menyoroti pergeseran dan perubahan tertentu dari waktu ke waktu.",
    "Statistik bergerak yang dihitung menunjukkan adanya tren spesifik.",
    "Penggunaan rata-rata bergerak memungkinkan identifikasi tren dalam kumpulan data.",
    "Tren yang diungkapkan oleh perhitungan rata-rata bergerak sangat signifikan."
]

time_range_context_id = [
    "Periode dari {start} hingga {end} sangat penting untuk memahami perubahan langsung dan pola yang lebih luas.",
    "Menganalisis data antara {start} dan {end} memberikan wawasan utama tentang variasi sementara dan jangka panjang.",
    "Memeriksa data yang mencakup {start} hingga {end} memungkinkan kita untuk mengamati variabilitas jangka pendek dan tren jangka panjang.",
    "Kerangka waktu dari {start} hingga {end} memberikan informasi penting tentang perilaku dinamis data.",
    "Dalam interval {start} hingga {end}, kita dapat melihat fluktuasi jangka pendek dan tren jangka panjang yang signifikan.",
    "Pengamatan dalam periode {start} hingga {end} sangat penting untuk memahami perubahan langsung dan berkelanjutan data.",
    "Memahami perilaku data selama periode dari {start} hingga {end} adalah kunci untuk mengidentifikasi perubahan langsung dan bertahap."
]

overall_insight_intro_id = [
    "Secara ringkas, analisis dengan jelas menunjukkan temuan spesifik.",
    "Sebagai kesimpulan, analisis memberikan bukti definitif dari kesimpulan tertentu.",
    "Pada dasarnya, analisis membuktikan poin-poin tertentu secara meyakinkan.",
    "Singkatnya, analisis memberikan hasil yang jelas dan final.",
    "Intinya, analisis data memberikan wawasan yang meyakinkan.",
    "Sederhananya, analisis mengarah pada kesimpulan yang jelas dan tegas.",
    "Pendeknya, analisis menunjukkan hasil yang pasti."
]

variables_detail_intro_id = [
    "Penilaian menyeluruh terhadap variabel ini menegaskan karakteristik spesifik.",
    "Studi mendalam tentang variabel ini memberikan konfirmasi atribut tertentu.",
    "Dengan melakukan evaluasi mendalam, kita dapat mengkonfirmasi properti spesifik variabel.",
    "Evaluasi terperinci dari variabel ini memvalidasi temuan spesifik.",
    "Analisis ekstensif dari variabel ini menegaskan aspek tertentu.",
    "Melalui tinjauan yang cermat, karakteristik variabel dikonfirmasi.",
    "Analisis terperinci mendukung kualitas spesifik dari variabel ini."
]

data_structure_context_id = [
    "Menggunakan struktur beragam dari data memungkinkan kita untuk menerapkan metode analitis spesifik secara efisien.",
    "Memanfaatkan struktur beragam dari kumpulan data memungkinkan penggunaan strategi analitis yang ditargetkan secara efektif.",
    "Keragaman struktural kumpulan data memfasilitasi penerapan teknik analitis yang sesuai.",
    "Dengan memanfaatkan variasi struktural data, kita dapat menerapkan pendekatan analitis khusus.",
    "Struktur beragam dari kumpulan data memungkinkan penerapan metode analitis spesifik secara efektif.",
    "Memanfaatkan kompleksitas struktural kumpulan data memungkinkan penerapan teknik analitis yang terfokus.",
    "Struktur data yang beragam mendukung penggunaan strategi analitis yang disesuaikan untuk berbagai jenis data."
]

observation_counts_summary_id = [
    "Menganalisis data berdasarkan segmen waktu menunjukkan bahwa jumlah pengamatan sangat bervariasi di berbagai periode.",
    "Segmentasi data berdasarkan waktu mengungkapkan perbedaan substansial dalam jumlah pengamatan antar periode.",
    "Segmentasi data berbasis waktu menyoroti variabilitas jumlah pengamatan dari waktu ke waktu.",
    "Dengan melakukan segmentasi data secara kronologis, kita mengamati fluktuasi yang cukup besar dalam jumlah pengamatan.",
    "Pembagian data berdasarkan interval waktu menunjukkan bahwa jumlah pengamatan sangat berbeda antar periode.",
    "Segmentasi temporal data mengungkapkan variasi signifikan dalam jumlah pengamatan di berbagai kerangka waktu.",
    "Ketika data disegmentasikan berdasarkan waktu, menjadi jelas bahwa jumlah pengamatan tidak konsisten di semua periode."
]

rolling_trend_summary_id = [
    "Analisis pergerakan mengkonfirmasi keberlanjutan tren selama periode yang diamati.",
    "Tren yang sedang berlangsung dikonfirmasi oleh analisis pergerakan sepanjang kerangka waktu yang diamati.",
    "Analisis menggunakan statistik pergerakan memvalidasi tren konsisten yang terlihat selama periode pengamatan.",
    "Perhitungan pergerakan menunjukkan bahwa tren tetap konsisten selama seluruh periode yang dianalisis.",
    "Tren data terbukti persisten melalui analisis pergerakan yang dilakukan pada periode yang diamati.",
    "Dengan menerapkan analisis pergerakan, keberlanjutan tren dikonfirmasi di seluruh data yang diamati.",
    "Tren konsisten selama periode yang diamati didukung oleh temuan analisis pergerakan."
]

time_series_freq_explanation_id = [
    "Berdasarkan analisis frekuensi, dapat disimpulkan bahwa data dikumpulkan dalam pola berulang yang konsisten.",
    "Inferensi frekuensi mengungkapkan bahwa data dikumpulkan sesuai dengan siklus periodik yang dapat diprediksi.",
    "Pola pengumpulan data yang teratur disarankan oleh analisis frekuensi yang dilakukan.",
    "Melalui inferensi frekuensi, kita dapat menyimpulkan bahwa data mengikuti proses pengumpulan siklus rutin.",
    "Bukti menunjukkan, berdasarkan analisis frekuensi, bahwa data dikumpulkan pada interval periodik yang teratur.",
    "Analisis frekuensi data menunjukkan siklus pengumpulan yang konsisten dan berulang.",
    "Jelas dari inferensi frekuensi bahwa data mematuhi jadwal pengumpulan periodik yang teratur."
]

final_overall_summary_id = [
    "Pada akhirnya, temuan gabungan menawarkan dasar yang kuat untuk strategi analitis masa depan dan pengembangan model.",
    "Sebagai kesimpulan, hasil yang dikumpulkan memberikan fondasi yang kuat untuk arah analisis data dan pembuatan model berikutnya.",
    "Untuk meringkas, temuan terintegrasi membangun platform yang kuat untuk kemajuan analisis data dan pembangunan model.",
    "Sebagai poin akhir, temuan yang dikonsolidasikan menciptakan dasar yang kokoh untuk masa depan analisis data dan pengembangan model.",
    "Sebagai penutup, temuan kolektif memberikan dasar yang stabil untuk jalur masa depan analisis data dan pembuatan model.",
    "Untuk menyimpulkan, temuan yang disintesis adalah titik awal yang dapat diandalkan untuk evolusi analisis data dan desain model.",
    "Akhirnya, hasil gabungan menyajikan dasar yang kuat untuk eksplorasi data masa depan dan konstruksi model."
]

mode_std_variance_summary_id = [
    "Dengan modus {mode_val}, data menunjukkan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}, mengindikasikan variabilitas yang signifikan dalam kumpulan data.",
    "Dengan {mode_val} sebagai nilai yang paling sering muncul, data menunjukkan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}, menyiratkan dispersi yang cukup besar.",
    "Data, dengan modus {mode_val}, menunjukkan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}, yang menunjukkan tingkat variabilitas yang signifikan.",
    "Modus {mode_val} diamati, dan standar deviasi kumpulan data adalah {std_val:.4f} dan varians adalah {var_val:.4f}, menyoroti variabilitas substansial.",
    "Memiliki modus {mode_val}, data mengungkapkan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}, menunjukkan sejumlah besar variasi.",
    "Nilai data yang paling umum adalah {mode_val}, dan menunjukkan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}, menunjukkan variabilitas yang signifikan.",
    "Dengan {mode_val} sebagai modus, data memiliki standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}, menunjukkan tingkat variabilitas yang tinggi."
]

range_iqr_summary_id = [
    "Analisis data mengungkapkan rentang keseluruhan sebesar {range_val:.4f} dan rentang interkuartil (IQR) sebesar {iqr_val:.4f}, yang keduanya secara kolektif menggambarkan penyebaran data.",
    "Analisis data menunjukkan rentang total sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}, menekankan dispersi kumpulan data.",
    "Analisis data menunjukkan bahwa kumpulan data memiliki rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}, yang bersama-sama menunjukkan penyebaran data.",
    "Rentang keseluruhan sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f} diamati dalam analisis data, menyoroti distribusi kumpulan data.",
    "Penyebaran data dibuktikan dengan rentangnya sebesar {range_val:.4f} dan IQR-nya sebesar {iqr_val:.4f}, seperti yang ditunjukkan oleh analisis.",
    "Dispersi kumpulan data ditandai dengan rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}, menurut analisis data.",
    "Analisis data menunjukkan penyebaran seperti yang ditunjukkan oleh rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}."
]

overall_dispersion_summary_id = [
    "Secara umum, dispersi data yang cukup besar disoroti oleh nilai standar deviasi yang tinggi ({std_val:.4f}) dan varians ({var_val:.4f}), bersama dengan rentang sebesar {range_val:.4f}.",
    "Penyebaran data yang signifikan terbukti dari standar deviasi yang tinggi ({std_val:.4f}) dan varians ({var_val:.4f}), serta rentang sebesar {range_val:.4f}.",
    "Secara keseluruhan, tingkat dispersi yang substansial ditunjukkan oleh standar deviasi yang besar ({std_val:.4f}) dan varians ({var_val:.4f}), bersama dengan rentang sebesar {range_val:.4f}.",
    "Dispersi kumpulan data ditunjukkan dengan jelas oleh standar deviasi yang tinggi ({std_val:.4f}) dan varians ({var_val:.4f}), bersama dengan rentang sebesar {range_val:.4f}.",
    "Dispersi yang luar biasa ditunjukkan oleh standar deviasi yang tinggi ({std_val:.4f}) dan varians ({var_val:.4f}), serta rentang sebesar {range_val:.4f} yang diamati dalam data.",
    "Dispersi data yang signifikan tercermin dalam standar deviasi yang tinggi ({std_val:.4f}), varians ({var_val:.4f}), dan rentang sebesar {range_val:.4f}.",
    "Nilai tinggi untuk standar deviasi ({std_val:.4f}) dan varians ({var_val:.4f}), ditambah dengan rentang sebesar {range_val:.4f}, secara kolektif menandakan dispersi data yang cukup besar."
]

distribution_metrics_summary_id = [
    "Analisis deskriptif menyoroti bahwa data memiliki rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, di samping rentang sebesar {range_val:.4f} dan rentang interkuartil (IQR) sebesar {iqr_val:.4f}.",
    "Dari analisis deskriptif, kita mengamati rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, dengan penyebaran data diwakili oleh rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}.",
    "Rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val} diungkapkan oleh analisis deskriptif, dengan dispersi data diukur dengan rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}.",
    "Analisis deskriptif menunjukkan bahwa tendensi sentral kumpulan data mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, sementara variabilitasnya ditangkap oleh rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}.",
    "Analisis statistik deskriptif menunjukkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, dengan penyebaran data didefinisikan oleh rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}.",
    "Metrik deskriptif menunjukkan rata-rata data menjadi {mean_val:.4f}, mediannya {median_val:.4f}, dan modusnya {mode_val}, dan penyebarannya lebih lanjut didefinisikan oleh rentang sebesar {range_val:.4f} dan IQR sebesar {iqr_val:.4f}.",
    "Karakteristik distribusi data, seperti yang diungkapkan oleh analisis deskriptif, mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, rentang sebesar {range_val:.4f}, dan IQR sebesar {iqr_val:.4f}."
]

central_tendency_dispersion_id = [
    "Distribusi keseluruhan digambarkan oleh ukuran pusat seperti modus sebesar {mode_val} dan median sebesar {median_val:.4f}, bersama dengan ukuran dispersi seperti standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}.",
    "Ukuran pusat, termasuk modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dikombinasikan dengan ukuran penyebaran seperti standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, menggambarkan distribusi keseluruhan data.",
    "Distribusi data ditandai oleh modus sebesar {mode_val} dan median sebesar {median_val:.4f}, yang mencerminkan tendensi sentral, dan standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, yang menunjukkan dispersi.",
    "Modus sebesar {mode_val} dan median sebesar {median_val:.4f} memberikan wawasan tentang tendensi sentral data, sementara standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f} menggambarkan dispersinya, bersama-sama menunjukkan distribusi keseluruhan.",
    "Distribusi keseluruhan diwakili oleh modus sebesar {mode_val} dan median sebesar {median_val:.4f}, menyoroti tendensi sentral, dan standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, menunjukkan penyebaran data.",
    "Tendensi sentral, seperti yang ditunjukkan oleh modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan dispersi, diwakili oleh standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, bersama-sama menentukan distribusi data.",
    "Distribusi keseluruhan data ditangkap oleh ukuran pusatnya, modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan ukuran dispersinya, standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}."
]

descriptive_stats_template_id = [
    "Rata-rata dan median, diukur masing-masing sebesar {mean_val:.4f} dan {median_val:.4f}, menunjukkan distribusi {distribution} dalam data.",
    "Rata-rata sebesar {mean_val:.4f} dan median sebesar {median_val:.4f} diamati, menunjukkan bahwa data mengikuti pola distribusi {distribution}.",
    "Metrik deskriptif mengungkapkan rata-rata sebesar {mean_val:.4f} dan median sebesar {median_val:.4f}, yang bersama-sama menunjukkan distribusi {distribution}.",
    "Mengingat rata-rata sebesar {mean_val:.4f} dan median sebesar {median_val:.4f}, distribusi data diklasifikasikan sebagai {distribution}.",
    "Data menunjukkan rata-rata sebesar {mean_val:.4f} dan median sebesar {median_val:.4f}, menunjukkan bahwa distribusinya dapat dikarakterisasi sebagai {distribution}.",
    "Dengan rata-rata sebesar {mean_val:.4f} dan median sebesar {median_val:.4f}, distribusi data menunjukkan pola {distribution}.",
    "Ukuran statistik deskriptif, rata-rata sebesar {mean_val:.4f} dan median sebesar {median_val:.4f}, menunjukkan distribusi {distribution}."
]

variability_insight_id = [
    "Dispersi data signifikan, dengan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f} menunjukkan tingkat penyebaran yang tinggi.",
    "Tingkat variabilitas yang signifikan diamati dalam data, sebagaimana dibuktikan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Data menunjukkan penyebaran yang cukup besar, disorot oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Penyebaran data tinggi ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Penyebaran data yang tinggi diukur dengan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Variabilitas signifikan dalam data ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Dispersi data substansial, seperti yang ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}."
]

statistical_measure_overview_id = [
    "Tinjauan ukuran statistik menunjukkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan distribusi nilai yang luas.",
    "Penilaian statistik mengungkapkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan bahwa data tersebar luas.",
    "Distribusi nilai yang luas disarankan oleh tinjauan statistik, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Nilai data terdistribusi luas, seperti yang ditunjukkan oleh penilaian statistik, yang melaporkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Secara keseluruhan, ukuran statistik, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan penyebaran yang luas dalam data.",
    "Distribusi data luas, seperti yang ditunjukkan oleh tinjauan statistik, yang memberikan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Distribusi nilai yang luas terbukti dari tinjauan ukuran statistik, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}."
]

statistical_measure_overview_id = [
    "Penilaian statistik keseluruhan mengungkapkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan bahwa nilai data tersebar luas.",
    "Distribusi nilai yang luas disarankan oleh tinjauan statistik, yang melaporkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Penyebaran data terbukti dari penilaian statistik, yang menunjukkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Pemeriksaan statistik keseluruhan, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan penyebaran yang luas dalam data.",
    "Tinjauan statistik menyoroti rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, yang bersama-sama menunjukkan distribusi nilai yang luas.",
    "Rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, seperti yang ditunjukkan dalam penilaian statistik, menunjukkan bahwa nilai data tersebar.",
    "Distribusi data yang luas dikonfirmasi oleh tinjauan statistik, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}."
]

central_tendency_dispersion_id = [
    "Distribusi keseluruhan divisualisasikan melalui ukuran pusat, seperti modus sebesar {mode_val} dan median sebesar {median_val:.4f}, bersama dengan ukuran dispersi, termasuk standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}.",
    "Tendensi sentral, ditunjukkan oleh modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan penyebaran data, ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, bersama-sama memberikan gambaran distribusi keseluruhan.",
    "Distribusi data ditandai oleh ukuran pusat, modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan ukuran variabilitas, standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}.",
    "Pandangan komprehensif tentang distribusi data ditawarkan oleh ukuran pusat, modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan ukuran dispersi, standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}.",
    "Distribusi data diwakili oleh tendensi sentralnya, modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan penyebarannya, standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}.",
    "Ukuran pusat, termasuk modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan ukuran dispersi, seperti standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, secara bersama-sama menggambarkan distribusi keseluruhan data.",
    "Distribusi data dijelaskan secara menyeluruh oleh ukuran pusatnya, modus sebesar {mode_val} dan median sebesar {median_val:.4f}, dan variabilitasnya, standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}."
]

variability_insight_id = [
    "Dispersi data signifikan, dengan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f} menunjukkan tingkat penyebaran yang tinggi.",
    "Tingkat variabilitas yang signifikan diamati dalam data, sebagaimana dibuktikan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Data menunjukkan penyebaran yang cukup besar, disorot oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Penyebaran data tinggi ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Penyebaran data yang tinggi diukur dengan standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Variabilitas signifikan dalam data ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}.",
    "Dispersi data substansial, seperti yang ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan varians sebesar {var_val:.4f}."
]

statistical_measure_overview_id = [
    "Tinjauan metrik statistik penting mengungkapkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan penyebaran data yang luas.",
    "Ukuran statistik utama, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan penyebaran yang luas dalam data.",
    "Memeriksa ukuran statistik inti, rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menyoroti distribusi data yang luas.",
    "Penyebaran data yang luas terbukti dari tinjauan statistik utama, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Distribusi nilai yang luas ditunjukkan oleh tinjauan statistik, yang menyajikan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Tinjauan ukuran statistik kritis, rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan penyebaran yang luas dalam data.",
    "Dispersi data ditunjukkan dengan jelas dalam tinjauan statistik, yang melaporkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}."
]

spread_and_skew_summary_id = [
    "Variabilitas dan asimetri distribusi ditunjukkan oleh standar deviasi terukur sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f}.",
    "Standar deviasi sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f} bersama-sama menggambarkan dispersi dan kurangnya simetri data.",
    "Penyebaran data dan penyimpangannya dari simetri ditangkap oleh standar deviasi sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f}.",
    "Variabilitas dan asimetri yang terukur dalam distribusi ditunjukkan oleh standar deviasi sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f}.",
    "Karakteristik distribusi, dalam hal penyebaran dan skew, didefinisikan oleh standar deviasi sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f}.",
    "Standar deviasi sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f} memberikan wawasan tentang dispersi dan asimetri data.",
    "Variabilitas dan asimetri distribusi data diindikasikan dengan jelas oleh standar deviasi sebesar {std_val:.4f} dan nilai skewness sebesar {skew_val:.4f}."
]

outlier_context_summary_id = [
    "Kehadiran {outlier_count} outlier yang teridentifikasi menunjukkan perlunya investigasi lebih lanjut untuk memastikan integritas data.",
    "Jumlah outlier sebanyak {outlier_count} menunjukkan bahwa analisis tambahan diperlukan untuk menilai keandalan data.",
    "Dengan {outlier_count} outlier terdeteksi, disarankan untuk melakukan investigasi lebih lanjut untuk memverifikasi akurasi data.",
    "Deteksi {outlier_count} outlier menyiratkan bahwa pengawasan lebih lanjut diperlukan untuk mengkonfirmasi kepercayaan data.",
    "Outlier {outlier_count} yang teridentifikasi memerlukan investigasi tambahan untuk mengevaluasi integritas data.",
    "Jumlah outlier sebanyak {outlier_count} memerlukan penilaian lebih lanjut untuk menentukan keandalan data.",
    "Kehadiran {outlier_count} outlier menunjukkan bahwa investigasi yang lebih menyeluruh diperlukan untuk memastikan validitas data."
]

variance_range_iqr_detail_id = [
    "Penyebaran keseluruhan data secara efektif ditangkap oleh varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan rentang interkuartil (IQR) sebesar {iqr_val:.4f}, yang semuanya menunjukkan dispersi yang luas.",
    "Dispersi luas ditunjukkan oleh varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan IQR sebesar {iqr_val:.4f}, yang semuanya mewakili penyebaran keseluruhan data.",
    "Varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan IQR sebesar {iqr_val:.4f} secara kolektif menggambarkan dispersi data yang luas dan penyebaran keseluruhan.",
    "Penyebaran keseluruhan data, seperti yang ditunjukkan oleh varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan IQR sebesar {iqr_val:.4f}, menunjukkan dispersi yang luas.",
    "Dispersi data secara komprehensif diwakili oleh varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan IQR sebesar {iqr_val:.4f}, yang semuanya menyoroti penyebaran keseluruhan data.",
    "Dispersi data yang luas ditunjukkan oleh varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan IQR sebesar {iqr_val:.4f}, yang bersama-sama menangkap penyebaran keseluruhannya.",
    "Penyebaran data keseluruhan secara menyeluruh diwakili oleh varians sebesar {var_val:.4f}, rentang sebesar {range_val:.4f}, dan rentang interkuartil sebesar {iqr_val:.4f}, menunjukkan dispersi yang luas."
]

mode_and_dispersion_detail_id = [
    "Nilai yang paling sering muncul dalam data adalah {mode_val}, dan ini, dikombinasikan dengan standar deviasi sebesar {std_val:.4f}, menunjukkan penyebaran yang signifikan dalam data.",
    "Dengan {mode_val} sebagai modus, data menunjukkan penyebaran yang cukup besar, sebagaimana dibuktikan oleh standar deviasi sebesar {std_val:.4f}.",
    "Modus data, {mode_val}, bila dipertimbangkan bersama dengan standar deviasi sebesar {std_val:.4f}, mengungkapkan tingkat dispersi yang signifikan.",
    "Standar deviasi sebesar {std_val:.4f} menunjukkan penyebaran data yang signifikan, terutama ketika modusnya adalah {mode_val}.",
    "Penyebaran data substansial, seperti yang ditunjukkan oleh standar deviasi sebesar {std_val:.4f}, mengingat nilai yang paling umum adalah {mode_val}.",
    "Modus, {mode_val}, bersama dengan standar deviasi sebesar {std_val:.4f}, menyoroti penyebaran data yang cukup besar.",
    "Dispersi signifikan dalam data disarankan oleh standar deviasi sebesar {std_val:.4f}, terutama ketika modusnya adalah {mode_val}."
]

statistical_overview_paraphrase_id = [
    "Penilaian statistik keseluruhan mengungkapkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan bahwa nilai data tersebar luas.",
    "Distribusi nilai yang luas disarankan oleh tinjauan statistik, yang melaporkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Penyebaran data terbukti dari penilaian statistik, yang menunjukkan rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Pemeriksaan statistik keseluruhan, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, menunjukkan penyebaran yang luas dalam data.",
    "Tinjauan statistik menyoroti rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, yang bersama-sama menunjukkan distribusi nilai yang luas.",
    "Rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, seperti yang ditunjukkan dalam penilaian statistik, menunjukkan bahwa nilai data tersebar.",
    "Distribusi data yang luas dikonfirmasi oleh tinjauan statistik, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}."
]

univariate_overview_metrics_id = [
    "Analisis univariat menyajikan fitur statistik utama variabel, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}.",
    "Ringkasan karakteristik statistik variabel, seperti rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, disediakan oleh analisis univariat.",
    "Analisis univariat menghasilkan ukuran statistik fundamental, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, secara efektif meringkas variabel.",
    "Metrik statistik utama, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, diperoleh melalui analisis univariat, menawarkan wawasan tentang perilaku variabel.",
    "Analisis univariat menghasilkan serangkaian nilai statistik inti—rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}—yang secara ringkas menggambarkan variabel.",
    "Properti statistik esensial variabel dienkapsulasi dalam rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, seperti yang ditentukan oleh analisis univariat.",
    "Ringkasan komprehensif statistik variabel, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, dicapai melalui analisis univariat."
]

detailed_statistical_insight_id = [
    "Profil statistik variabel, termasuk rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, dan modus sebesar {mode_val}, bersama dengan standar deviasi sebesar {std_val:.4f} dan rentang sebesar {range_val:.4f}, dikonfirmasi secara menyeluruh oleh analisis statistik terperinci.",
    "Analisis statistik terperinci memberikan konfirmasi tendensi sentral dan dispersi variabel, seperti yang diwakili oleh rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, standar deviasi sebesar {std_val:.4f}, dan rentang sebesar {range_val:.4f}.",
    "Karakteristik distribusi variabel, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, standar deviasi sebesar {std_val:.4f}, dan rentang sebesar {range_val:.4f}, divalidasi oleh analisis statistik terperinci.",
    "Gambaran komprehensif properti statistik variabel, termasuk tendensi sentral dan dispersinya, ditawarkan oleh analisis terperinci, yang mengkonfirmasi rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, standar deviasi sebesar {std_val:.4f}, dan rentang sebesar {range_val:.4f}.",
    "Analisis statistik terperinci mendukung fitur statistik variabel, yang mencakup rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, standar deviasi sebesar {std_val:.4f}, dan rentang sebesar {range_val:.4f}.",
    "Susunan statistik variabel, seperti yang didefinisikan oleh rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, standar deviasi sebesar {std_val:.4f}, dan rentang sebesar {range_val:.4f}, dikonfirmasi melalui analisis statistik terperinci.",
    "Melalui analisis statistik terperinci, variabel ditunjukkan memiliki rata-rata sebesar {mean_val:.4f}, median sebesar {median_val:.4f}, modus sebesar {mode_val}, standar deviasi sebesar {std_val:.4f}, dan rentang sebesar {range_val:.4f}, yang semuanya secara kolektif menggambarkan distribusinya."
]

# 1. Pendahuluan Analisis Deret Waktu (untuk analisis rentang waktu)
time_series_analysis_intro_id = [
    "Komponen analisis deret waktu mencakup periode dari {start_date} hingga {end_date}. Rentang waktu yang panjang ini memungkinkan kami menangkap tren jangka panjang, siklus halus, dan pergeseran mendadak yang krusial untuk peramalan.",
    "Mencakup periode dari {start_date} hingga {end_date}, dataset ini memberikan konteks temporal yang komprehensif, yang memungkinkan deteksi tren yang konsisten, pola musiman berulang, dan anomali sporadis.",
    "Periode analisis yang berlangsung dari {start_date} hingga {end_date} menyediakan kerangka waktu yang luas, sehingga kami dapat mengidentifikasi tren tahan lama serta variasi jangka pendek yang signifikan.",
    "Dari {start_date} hingga {end_date}, dataset ini menawarkan landasan temporal yang kuat, membantu mengungkap tren bertahap, fluktuasi musiman, serta gangguan sesaat yang dapat mempengaruhi keputusan strategis.",
    "Rentang pengamatan dari {start_date} hingga {end_date} memberikan dasar yang kokoh untuk analisis deret waktu, memungkinkan penangkapan evolusi data secara mendalam serta identifikasi siklus dan anomali.",
    "Analisis data dari {start_date} hingga {end_date} mengungkap konteks temporal yang kaya, sehingga memungkinkan identifikasi tren berkelanjutan, siklus musiman, dan deviasi sesaat yang informatif."
]

# 2. Insight Pola Musiman (untuk masing-masing variabel numerik yang dikelompokkan berdasarkan frekuensi yang dipilih)
seasonal_pattern_insight_id = [
    "Untuk '{col}', hasil resampling berdasarkan frekuensi {timefreq} menunjukkan bahwa rata-rata tertinggi adalah {max_val:.4f} pada periode {max_period} (yang {max_status} rata-rata keseluruhan),\ndan rata-rata terendah adalah {min_val:.4f} pada periode {min_period} (yang {min_status} rata-rata keseluruhan).\n\nHal ini menunjukkan adanya siklus musiman yang konsisten pada variabel tersebut.",
    "Analisis musiman untuk '{col}' mengungkap bahwa saat data dikelompokkan berdasarkan frekuensi {timefreq}, periode {max_period} mencatat rata-rata tertinggi sebesar {max_val:.4f} ({max_status} rata-rata keseluruhan),\ndan periode {min_period} mencatat rata-rata terendah sebesar {min_val:.4f} ({min_status} rata-rata keseluruhan).\n\nTemuan ini menandakan adanya fluktuasi musiman yang signifikan.",
    "Pada variabel '{col}', pengelompokan berdasarkan frekuensi {timefreq} menunjukkan bahwa nilai rata-rata maksimum sebesar {max_val:.4f} terjadi pada periode {max_period} ({max_status} rata-rata keseluruhan),\ndan nilai rata-rata minimum sebesar {min_val:.4f} terjadi pada periode {min_period} ({min_status} rata-rata keseluruhan).\n\nPola ini mengindikasikan adanya variasi musiman yang nyata.",
    "Hasil resampling untuk '{col}' dengan frekuensi {timefreq} mengungkap bahwa puncak rata-rata terjadi pada periode {max_period} dengan nilai {max_val:.4f} ({max_status} rata-rata keseluruhan),\ndan lembah rata-rata terjadi pada periode {min_period} dengan nilai {min_val:.4f} ({min_status} rata-rata keseluruhan).\n\nAnalisis ini memperlihatkan adanya siklus musiman yang jelas.",
    "Untuk '{col}', saat data dikelompokkan berdasarkan frekuensi {timefreq}, periode {max_period} menunjukkan nilai rata-rata tertinggi ({max_val:.4f}, {max_status} rata-rata keseluruhan),\ndan periode {min_period} menunjukkan nilai rata-rata terendah ({min_val:.4f}, {min_status} rata-rata keseluruhan).\n\nTemuan ini menegaskan adanya pola musiman dalam variabel tersebut.",
    "Analisis musiman pada '{col}' berdasarkan frekuensi {timefreq} menunjukkan bahwa periode dengan rata-rata tertinggi adalah {max_period} (nilai: {max_val:.4f}, {max_status} rata-rata keseluruhan),\ndan periode dengan rata-rata terendah adalah {min_period} (nilai: {min_val:.4f}, {min_status} rata-rata keseluruhan).\n\nPola ini mengindikasikan variasi musiman yang signifikan."
]

# 3. Insight Statistik Bergulir dan Tren (untuk masing-masing variabel numerik)
rolling_trend_insight_id = [
    "Untuk '{col}', analisis rolling dengan jendela 5 periode menghasilkan rata-rata bergerak sebesar {roll_mean:.4f} dan standar deviasi bergerak sebesar {roll_std:.4f}.\n\nHal ini menunjukkan bahwa tren dasar pada variabel ini relatif stabil, meskipun terdapat fluktuasi periodik yang konsisten.",
    "Analisis rolling untuk '{col}' (menggunakan jendela 5 periode) mengungkapkan bahwa nilai rata-rata bergerak adalah {roll_mean:.4f} dengan standar deviasi bergerak sebesar {roll_std:.4f}.\n\nTemuan ini menunjukkan adanya tren yang stabil, namun dengan variasi berkala yang patut dicermati untuk memahami dinamika variabel.",
    "Evaluasi untuk '{col}' melalui rolling window 5 periode menghasilkan rata-rata bergerak {roll_mean:.4f} dan standar deviasi bergerak {roll_std:.4f}.\n\nHasil ini menyiratkan bahwa variabel mempertahankan tren konsisten dengan adanya variasi periodik yang dapat mengindikasikan pergeseran atau siklus dalam data.",
    "Pada variabel '{col}', perhitungan statistik rolling dengan jendela 5 periode menunjukkan rata-rata bergerak rata-rata {roll_mean:.4f} dan standar deviasi rata-rata {roll_std:.4f}.\n\nAnalisis ini menggambarkan tren yang cukup stabil dengan penyimpangan yang teratur, yang dapat mencerminkan siklus atau pergeseran dalam nilai-nilai variabel.",
    "Hasil analisis rolling pada '{col}' menggunakan jendela 5 periode menunjukkan bahwa rata-rata bergerak adalah {roll_mean:.4f} dan standar deviasi bergerak adalah {roll_std:.4f}.\n\nData ini mengindikasikan tren dasar yang stabil, namun terdapat fluktuasi periodik yang moderat, memberikan gambaran tentang dinamika variabel secara keseluruhan.",
    "Analisis statistik bergulir untuk '{col}' (jendela 5 periode) menghasilkan rata-rata bergerak sebesar {roll_mean:.4f} dan standar deviasi bergerak sebesar {roll_std:.4f}.\n\nTemuan ini menandakan bahwa meskipun variabel menunjukkan tren yang konsisten, terdapat variasi periodik yang signifikan yang patut diinvestigasi lebih lanjut."
]

# ============================================== DOC REPORTING ===========================================================================================================

def descriptive_freq(freq_code):
    """Map a pandas frequency code to a descriptive frequency string."""
    if not freq_code or freq_code == "":
        return "Not determined"
    freq_code = freq_code.upper()
    if freq_code in ["D", "B"]:
        return "Daily"
    if freq_code.startswith("W"):
        return "Weekly"
    if freq_code.startswith("M") or freq_code.startswith("MS") or freq_code.startswith("BM") or freq_code.startswith("CBM"):
        return "Monthly"
    if "Q" in freq_code:
        return "Quarterly"
    if freq_code.startswith("A") or freq_code.startswith("Y"):
        return "Yearly"
    if freq_code.startswith("H"):
        return "Hourly"
    return freq_code

# Define regex patterns for various date formats
date_patterns = [
    r'\d{4}-\d{2}-\d{2}',           # YYYY-MM-DD (e.g., 2025-03-13)
    r'\d{2}-\d{2}-\d{4}',           # DD-MM-YYYY (e.g., 13-03-2025)
    r'\d{4}/\d{2}/\d{2}',           # YYYY/MM/DD (e.g., 2025/03/13)
    r'\d{2}/\d{2}/\d{4}',           # DD/MM/YYYY (e.g., 13/03/2025)
    r'\d{2}/\d{2}/\d{4}',           # MM/DD/YYYY (e.g., 03/13/2025, context matters)
    r'\d{4}-\d{2}',               # YYYY-MM (e.g., 2025-03)
    r'\d{2}-\d{4}',               # MM-YYYY (e.g., 03-2025)
    r'\d{4}/\d{2}',               # YYYY/MM (e.g., 2025/03)
    r'\d{2}/\d{4}',               # MM/YYYY (e.g., 03/2025)
    r'\d{4}',                   # YYYY (e.g., 2025)
    r'\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}', # YYYY-MM-DD HH:MM:SS (e.g., 2025-03-13 10:30:45)
    r'\d{4}-\d{2}-\d{2} \d{2}:\d{2}',     # YYYY-MM-DD HH:MM (e.g., 2025-03-13 10:30)
    r'\d{2}-\d{2}-\d{4} \d{2}:\d{2}:\d{2}',   # DD-MM-YYYY HH:MM:SS (e.g., 13-03-2025 10:30:45)
    r'\d{8}',                       # YYYYMMDD (e.g., 20250313)
    r'\d{10}',                      # Unix Timestamp (e.g., 1647211200)

    # Quarterly Patterns
    r'\d{4}[-]?Q',              # YYYY[-]?Qq (e.g., 2025-Q1 or 2025Q1)
    r'Q[-]?\d{4}',               # Qq[-]?YYYY (e.g., Q1-2025 or Q12025)
    r'Q/\d{4}',              # Qq/YYYY (e.g., Q1/2025)

    # Yearly Patterns (already covered by '\d{4}')
    
    # Weekly Patterns
    r'\d{4}-W\d{2}',             # YYYY-Www (e.g., 2025-W05)
    r'\d{4}W\d{2}',              # YYYYWww (e.g., 2025W05)
    r'\d{4}-W\d{2}-\d{1}',       # YYYY-Www-D (e.g., 2025-W05-1) (Day of week)
    r'\d{4}W\d{2}-\d{1}',        # YYYYWww-D (e.g., 2025W05-1)
    r'\d{4}-W\d{2}\d{1}',        # YYYY-WwwD (e.g., 2025-W051)
    r'\d{4}W\d{2}\d{1}',         # YYYYWwwD (e.g., 2025W051)
    r'\d{4}-W\d{1,2}',           # YYYY-Ww or YYYY-WW
    r'\d{4}W{1,2}',            # YYYYWw or YYYYWW
    r'W\d{1,2}-\d{4}',          # Ww-YYYY or WW-YYYY
    r'W\d{1,2}/\d{4}',          # Ww/YYYY or WW/YYYY
    r'W\d{1,2}\d{4}',           # WwYYYY or WWYYYY
]

def is_date_string(s):
    """Returns True if the string s matches any of the allowed date formats."""
    for pattern in date_patterns:
        if re.fullmatch(pattern, s):
            return True
    return False

def generate_doc_report_en(df):
    document = Document()
    
    # --- Introduction and Basic Descriptive Statistics ---
    document.add_heading("Comprehensive Exploratory Data Analysis Report", level=1)
    document.add_paragraph(
        f"This report provides an in-depth analysis of a dataset with {df.shape[0]} rows and {df.shape[1]} columns. "
        "It covers basic descriptive statistics, data structure, univariate analysis, and bivariate correlation analysis. "
        "The insights obtained serve as a strong foundation for further data cleaning, transformation, and predictive modeling."
    )
    
    # --- Data Types and Structure ---
    document.add_heading("Data Types and Structure", level=2)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    # Optionally, exclude any datetime columns if present (if needed)
    datetime_cols = []  # Modify this if you want to exclude specific datetime columns
    numeric_cols = [col for col in numeric_cols if col not in datetime_cols]
    object_cols = df.select_dtypes(include='object').columns.tolist()
    bool_cols = df.select_dtypes(include='bool').columns.tolist()
    category_cols = df.select_dtypes(include='category').columns.tolist()
    document.add_paragraph(
        "The dataset features a diverse range of data types. Numeric variables (e.g., " +
        (", ".join(numeric_cols) if numeric_cols else "None") +
        ") provide essential quantitative insights, while categorical and boolean variables add qualitative context. "
        "This analysis leverages these statistics to present a comprehensive overview of the data structure."
    )
    
    # --- Univariate Analysis and Outlier Detection ---
    document.add_heading("Univariate Analysis and Outlier Detection", level=2)
    if numeric_cols:
        for col in numeric_cols:
            document.add_heading(f"Analysis of Numeric Variable: '{col}'", level=3)
            series = df[col].dropna()
            if series.empty:
                document.add_paragraph(f"No data available for '{col}' after excluding missing values.")
                continue
            mean_val = series.mean()
            median_val = series.median()
            try:
                mode_val = series.mode()[0]
            except Exception:
                mode_val = "Not Available"
            std_val = series.std()
            var_val = series.var()
            min_val = series.min()
            max_val = series.max()
            range_val = max_val - min_val
            q25 = series.quantile(0.25)
            q75 = series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = series.skew()
            kurt_val = series.kurt()
            lower_bound = q25 - 1.5 * iqr_val
            upper_bound = q75 + 1.5 * iqr_val
            outlier_count = series[(series < lower_bound) | (series > upper_bound)].count()
            distribution = "balanced" if abs(mean_val - median_val) < std_val * 0.1 else "skewed"
            document.add_paragraph(
                f"The variable '{col}' exhibits a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}. "
                f"It shows a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, resulting in an overall range of {range_val:.4f}. "
                f"The interquartile range (IQR) is {iqr_val:.4f}, with Q1 at {q25:.4f} and Q3 at {q75:.4f}. "
                f"Skewness is {skew_val:.4f} and kurtosis is {kurt_val:.4f}, with {outlier_count} outlier(s) detected, providing a robust picture of the distribution."
            )
    else:
        document.add_paragraph("No numeric variables are available for univariate analysis.")
    
    # --- Bivariate Analysis: Correlation Analysis ---
    document.add_heading("Bivariate Analysis: Correlation Analysis", level=2)
    if numeric_cols and len(numeric_cols) > 1:
        corr_matrix = df[numeric_cols].corr()
        for col in numeric_cols:
            document.add_heading(f"Correlation Analysis for '{col}'", level=3)
            corr_series = corr_matrix[col].drop(labels=[col]).sort_values(ascending=False)
            top_n = min(5, len(corr_series))
            if top_n > 0:
                top_corr = corr_series.iloc[:top_n]
                corr_text = ", ".join([f"{idx} ({corr:.4f})" for idx, corr in top_corr.items()])
                document.add_paragraph(
                    f"For '{col}', the top {top_n} positively correlated variables are: {corr_text}. "
                    "These correlations indicate potential interdependencies that may be important for predictive modeling."
                )
            else:
                document.add_paragraph(f"No significant correlations detected for '{col}'.")
    else:
        document.add_paragraph("Bivariate correlation analysis could not be performed due to insufficient numeric variables.")
    
    # --- Variables Insight and Overall Implications ---
    document.add_heading("Variables Insight and Overall Implications", level=2)
    overall_intro = random.choice(overall_insight_intro_en) if 'overall_insight_intro_en' in globals() else ""
    document.add_paragraph(
        overall_intro + " In this final section, each variable is examined in detail to extract actionable insights that serve as a basis for further analysis. "
        "By synthesizing key statistical measures and observed patterns, we develop a comprehensive understanding of the dataset's structure, variability, and anomalies, "
        "which can be leveraged for targeted data cleaning, transformation, and predictive modeling."
    )
    
    for col in df.columns:
        # Exclude any datetime-related columns if applicable
        if col in datetime_cols:
            continue
        document.add_heading(f"Detailed Insight for '{col}'", level=3)
        var_detail_intro = random.choice(variables_detail_intro_en) if 'variables_detail_intro_en' in globals() else ""
        document.add_paragraph(var_detail_intro)
        col_series = df[col]
        base_info = (
            f"Data type: {col_series.dtype}. This variable contains {col_series.nunique()} unique values and {col_series.isnull().sum()} missing entries. "
        )
        if pd.api.types.is_numeric_dtype(col_series):
            mean_val = col_series.mean()
            median_val = col_series.median()
            std_val = col_series.std()
            var_val = col_series.var()
            min_val = col_series.min()
            max_val = col_series.max()
            range_val = max_val - min_val
            q25 = col_series.quantile(0.25)
            q75 = col_series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = col_series.skew()
            kurt_val = col_series.kurt()
            distribution = "balanced" if abs(mean_val - median_val) < std_val * 0.1 else "skewed"
            descriptive_text = random.choice(descriptive_stats_template_en).format(
                mean_val=mean_val, median_val=median_val, mode_val="N/A", distribution=distribution
            ) if 'descriptive_stats_template_en' in globals() else f"The mean is {mean_val:.4f} and the median is {median_val:.4f}."
            variability_text = random.choice(variability_insight_en).format(std_val=std_val, var_val=var_val) if 'variability_insight_en' in globals() else ""
            detailed_info = (
                f"For this numeric variable, {descriptive_text} The standard deviation is {std_val:.4f} and the variance is {var_val:.4f}, indicating significant dispersion. {variability_text}"
            )
        elif pd.api.types.is_object_dtype(col_series) or pd.api.types.is_bool_dtype(col_series):
            try:
                mode_val = col_series.mode()[0]
                detailed_info = (
                    f"For this categorical variable, the most frequently occurring value (mode) is '{mode_val}', indicating a dominant category that could influence further analysis."
                )
            except Exception:
                detailed_info = "The mode could not be determined for this variable."
        else:
            detailed_info = "Detailed statistical analysis could not be performed for this variable type."
        document.add_paragraph(base_info + detailed_info)
    
    final_summary = random.choice(final_overall_summary_en) if 'final_overall_summary_en' in globals() else ""
    document.add_paragraph(
        "The comprehensive insights presented in this report form a robust foundation for further analytical work. "
        "Understanding each variable's behavior—from central tendencies and dispersion to anomalies—enables targeted data cleaning, transformation, and predictive modeling. "
        "Advanced visualizations and statistical techniques further enhance interpretability and support data-driven decision making. " + final_summary
    )
    
    return document

def generate_doc_report_en_ts(df):
    # Duplicate the chosen datetime column into a new column "date_val"
    if ts_choice in df.columns:
        df["date_val"] = df[ts_choice]
    
    # Explicitly convert the chosen datetime column to datetime objects using pd.to_datetime
    if ts_choice in df.columns:
        df[ts_choice] = pd.to_datetime(df[ts_choice], errors='raise')
        df.set_index(ts_choice, inplace=True)
    
    document = Document()
    
    # --- Introduction and Basic Descriptive Statistics ---
    document.add_heading("Comprehensive Exploratory Data Analysis Report", level=1)
    document.add_paragraph(
        f"This report provides an in-depth analysis of a dataset with {df.shape[0]} rows and {df.shape[1]} columns. "
        "The time series analysis is designed to reveal long-term trends, cyclic behaviors, and abrupt shifts over an extended observation period. "
        "These insights are crucial for accurate forecasting and informed decision-making."
    )
    
    # --- Data Types and Structure ---
    document.add_heading("Data Types and Structure", level=2)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    datetime_cols = [ts_choice, "date_val"]
    numeric_cols = [col for col in numeric_cols if col not in datetime_cols]
    object_cols = df.select_dtypes(include='object').columns.tolist()
    bool_cols = df.select_dtypes(include='bool').columns.tolist()
    category_cols = df.select_dtypes(include='category').columns.tolist()
    document.add_paragraph(
        "The dataset features a diverse range of data types. Numeric variables (for example, " +
        (", ".join(numeric_cols) if numeric_cols else "None") +
        ") provide essential quantitative insights, while categorical and boolean variables offer qualitative context. "
        f"The datetime column '{ts_choice}' has been converted to a proper datetime format and set as the index, with its original values preserved in 'date_val'."
    )
    
    # --- Univariate Analysis and Outlier Detection (excluding datetime and date_val) ---
    document.add_heading("Univariate Analysis and Outlier Detection", level=2)
    if numeric_cols:
        for col in numeric_cols:
            document.add_heading(f"Analysis of Numeric Variable: '{col}'", level=3)
            series = df[col].dropna()
            if series.empty:
                document.add_paragraph(f"No data available for '{col}' after excluding missing values.")
                continue
            mean_val = series.mean()
            median_val = series.median()
            try:
                mode_val = series.mode()[0]
            except Exception:
                mode_val = "Not Available"
            std_val = series.std()
            var_val = series.var()
            min_val = series.min()
            max_val = series.max()
            range_val = max_val - min_val
            q25 = series.quantile(0.25)
            q75 = series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = series.skew()
            kurt_val = series.kurt()
            lower_bound = q25 - 1.5 * iqr_val
            upper_bound = q75 + 1.5 * iqr_val
            outlier_count = series[(series < lower_bound) | (series > upper_bound)].count()
            distribution = "balanced" if abs(mean_val - median_val) < std_val * 0.1 else "skewed"
            document.add_paragraph(
                f"The variable '{col}' exhibits a mean of {mean_val:.4f}, a median of {median_val:.4f}, and a mode of {mode_val}. "
                f"It shows a standard deviation of {std_val:.4f} and a variance of {var_val:.4f}, resulting in an overall range of {range_val:.4f}. "
                f"The interquartile range (IQR) is {iqr_val:.4f}, with Q1 at {q25:.4f} and Q3 at {q75:.4f}. "
                f"Skewness is {skew_val:.4f} and kurtosis is {kurt_val:.4f}, with {outlier_count} outlier(s) detected, providing a robust picture of the variable's distribution."
            )
    else:
        document.add_paragraph("No numeric variables are available for univariate analysis.")
    
    # --- Bivariate Analysis: Correlation Analysis ---
    document.add_heading("Bivariate Analysis: Correlation Analysis", level=2)
    if numeric_cols and len(numeric_cols) > 1:
        corr_matrix = df[numeric_cols].corr()
        for col in numeric_cols:
            document.add_heading(f"Correlation Analysis for '{col}'", level=3)
            corr_series = corr_matrix[col].drop(labels=[col]).sort_values(ascending=False)
            top_n = min(5, len(corr_series))
            if top_n > 0:
                top_corr = corr_series.iloc[:top_n]
                corr_text = ", ".join([f"{idx} ({corr:.4f})" for idx, corr in top_corr.items()])
                document.add_paragraph(
                    f"For '{col}', the top {top_n} positively correlated variables are: {corr_text}. "
                    "These correlations indicate potential interdependencies important for predictive modeling."
                )
            else:
                document.add_paragraph(f"No significant correlations detected for '{col}'.")
    else:
        document.add_paragraph("Bivariate correlation analysis could not be performed due to insufficient numeric variables.")
    
    # --- Time Series Analysis ---
    document.add_heading("Time Series Analysis", level=2)
    # Use the auxiliary "date_val" column to display the original datetime values
    if "date_val" in df.columns:
        start_date = df["date_val"].iloc[0]
        end_date = df["date_val"].iloc[-1]
    else:
        start_date = df.index[0]
        end_date = df.index[-1]
    ts_intro = random.choice(time_series_analysis_intro_en) if 'time_series_analysis_intro_en' in globals() else (
        "The time series spans from {start_date} to {end_date}, providing a robust temporal framework for analysis."
    )
    document.add_paragraph(ts_intro.format(start_date=start_date, end_date=end_date))
    
    # --- Seasonal Pattern Analysis using User-Defined Frequency ---
    document.add_heading("Seasonal Pattern Analysis", level=3)
    freq_mapping = {"Seconds": "S", "Minutes": "T", "Hours": "H", "Days": "D", "Weeks": "W", "Months": "M", "Quarterly": "Q", "Yearly": "Y"}
    freq_code = freq_mapping.get(timefreq, "D")
    season_text_full = (
        f"By resampling the data using the selected frequency ({timefreq}), we can examine the seasonal behavior of each numeric variable. "
        "The analysis computes the overall mean, standard deviation, and quartile measures (first quartile, median, and third quartile) for each resampled period. "
        "When these statistical measures are closely grouped, it suggests stable seasonal behavior; larger deviations indicate higher variability and more pronounced seasonal cycles. "
        "This comprehensive evaluation provides deep insights into the periodic behavior of the variables. \n\n"
    )
    for col in numeric_cols:
        try:
            resampled = df[col].resample(freq_code).mean()
            if resampled.count() < 2:
                season_text_full += f"For '{col}', there is insufficient data within each {timefreq.lower()} to compute meaningful seasonal statistics.\n\n"
                continue
            overall_avg = resampled.mean()
            overall_std = resampled.std()
            q1 = resampled.quantile(0.25)
            median = resampled.quantile(0.50)
            q3 = resampled.quantile(0.75)
            season_text_full += (
                f"For '{col}':\n"
                f"In the resampled data, the overall mean is {overall_avg:.4f} with a standard deviation of {overall_std:.4f}. "
                f"The first quartile (Q1) is {q1:.4f}, the median is {median:.4f}, and the third quartile (Q3) is {q3:.4f}. "
                "This indicates that if the quartile values are tightly clustered around the mean, the seasonal behavior is consistent; "
                "if they are widely spread, it suggests significant variability and more dynamic seasonal swings. \n\n"
            )
        except Exception as e:
            season_text_full += f"Seasonal analysis for '{col}' could not be completed due to error: {e}.\n\n"
    document.add_paragraph(season_text_full)
    
    # --- Rolling Statistics and Trend Analysis ---
    document.add_heading("Rolling Statistics and Trend Analysis", level=3)
    roll_text_full = ""
    for col in numeric_cols:
        try:
            window_results = []
            for window in range(1, 13):
                roll_mean_val = df[col].rolling(window=window, min_periods=1).mean().mean()
                roll_std_val = df[col].rolling(window=window, min_periods=1).std().mean()
                window_results.append((window, roll_mean_val, roll_std_val))
            best_mean = max(window_results, key=lambda x: x[1])
            worst_mean = min(window_results, key=lambda x: x[1])
            highest_std = max([r for r in window_results if not np.isnan(r[2])], key=lambda x: x[2], default=(None, None, None))
            lowest_std = min([r for r in window_results if not np.isnan(r[2])], key=lambda x: x[2], default=(None, None, None))
            roll_text_full += (
                f"For '{col}':\n"
                f"An exploration across window sizes from 1 to 12 {timefreq.lower()}(s) revealed that the highest average rolling mean of {best_mean[1]:.4f} was achieved with a window size of {best_mean[0]} {timefreq.lower()}(s), while the lowest average rolling mean of {worst_mean[1]:.4f} occurred with a window size of {worst_mean[0]} {timefreq.lower()}(s). "
                f"In terms of variability, the maximum rolling standard deviation, which indicates higher volatility, was {highest_std[2]:.4f} using a window size of {highest_std[0]} {timefreq.lower()}(s), and the minimum rolling standard deviation, indicating greater stability, was {lowest_std[2]:.4f} with a window size of {lowest_std[0]} {timefreq.lower()}(s). "
                "These insights help us understand how the variable behaves over different aggregation windows, highlighting both short-term fluctuations and long-term trends. \n\n"
            )
        except Exception as e:
            roll_text_full += f"Rolling statistics for '{col}' could not be computed due to error: {e}.\n\n"
    document.add_paragraph(roll_text_full)
    
    # --- Advanced Time Series Pattern Analysis ---
    document.add_heading("Advanced Time Series Pattern Analysis", level=3)
    adv_text_full = ""
    for col in numeric_cols:
        try:
            if df[col].dropna().empty:
                adv_text_full += f"No data available for advanced analysis of '{col}'.\n\n"
                continue
            # Use the auxiliary column 'date_val' to extract the original datetime values for extreme observations
            if "date_val" in df.columns:
                # If the index of the max/min is an integer, then use it as row number; otherwise, use it directly
                max_idx = df[col].idxmax()
                min_idx = df[col].idxmin()
                # Look up the corresponding date_val if available
                if isinstance(max_idx, int) and max_idx < len(df):
                    max_date = df["date_val"].iloc[max_idx]
                else:
                    max_date = max_idx
                if isinstance(min_idx, int) and min_idx < len(df):
                    min_date = df["date_val"].iloc[min_idx]
                else:
                    min_date = min_idx
            else:
                max_date = df[col].idxmax()
                min_date = df[col].idxmin()
            max_val = df[col].max()
            min_val = df[col].min()
            ts_detail = (
                f"For '{col}', the highest observed value is {max_val:.4f} recorded on {max_date}, and the lowest observed value is {min_val:.4f} recorded on {min_date}. "
            )
            try:
                adf_result = adfuller(df[col].dropna())
                p_value = adf_result[1]
                stationarity = "stationary" if p_value < 0.05 else "non-stationary"
                ts_detail += f"An extensive Augmented Dickey-Fuller (ADF) test indicates the series is {stationarity} with a p-value of {p_value:.4f}."
            except Exception as adf_err:
                ts_detail += f"ADF test could not be performed: {adf_err}."
            adv_text_full += ts_detail + "\n\n"
        except Exception as e:
            adv_text_full += f"Advanced time series analysis for '{col}' encountered an error: {e}.\n\n"
    document.add_paragraph(adv_text_full)
    
    # --- Variables Insight and Overall Implications ---
    document.add_heading("Variables Insight and Overall Implications", level=2)
    overall_intro = random.choice(overall_insight_intro_en) if 'overall_insight_intro_en' in globals() else ""
    document.add_paragraph(
        overall_intro + " In this final section, each variable is examined in detail to extract actionable insights that guide further analysis. "
        "By synthesizing key statistics and observed patterns, we develop a comprehensive understanding of the dataset’s structure, variability, and anomalies, "
        "forming a solid foundation for targeted data cleaning, transformation, and predictive modeling."
    )
    
    for col in df.columns:
        if col == ts_choice or col == "date_val":
            continue
        document.add_heading(f"Detailed Insight for '{col}'", level=3)
        var_detail_intro = random.choice(variables_detail_intro_en) if 'variables_detail_intro_en' in globals() else ""
        document.add_paragraph(var_detail_intro)
        col_series = df[col]
        base_info = (
            f"Data type: {col_series.dtype}. This variable contains {col_series.nunique()} unique values and {col_series.isnull().sum()} missing entries. "
        )
        if pd.api.types.is_numeric_dtype(col_series):
            mean_val = col_series.mean()
            median_val = col_series.median()
            std_val = col_series.std()
            var_val = col_series.var()
            min_val = col_series.min()
            max_val = col_series.max()
            range_val = max_val - min_val
            q25 = col_series.quantile(0.25)
            q75 = col_series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = col_series.skew()
            kurt_val = col_series.kurt()
            distribution = "balanced" if abs(mean_val - median_val) < std_val * 0.1 else "skewed"
            descriptive_text = random.choice(descriptive_stats_template_en).format(
                mean_val=mean_val, median_val=median_val, mode_val="N/A", distribution=distribution
            ) if 'descriptive_stats_template_en' in globals() else f"The mean is {mean_val:.4f} and the median is {median_val:.4f}."
            variability_text = random.choice(variability_insight_en).format(std_val=std_val, var_val=var_val) if 'variability_insight_en' in globals() else ""
            detailed_info = (
                f"For this numeric variable, {descriptive_text} The standard deviation is {std_val:.4f} and the variance is {var_val:.4f}, indicating significant dispersion. {variability_text}"
            )
        elif pd.api.types.is_object_dtype(col_series) or pd.api.types.is_bool_dtype(col_series):
            try:
                mode_val = col_series.mode()[0]
                detailed_info = (
                    f"For this categorical variable, the most frequently occurring value (mode) is '{mode_val}', indicating a dominant category that could influence further analysis."
                )
            except Exception:
                detailed_info = "The mode could not be determined for this variable."
        else:
            detailed_info = "Detailed statistical analysis could not be performed for this variable type."
        document.add_paragraph(base_info + detailed_info)
    
    final_summary = random.choice(final_overall_summary_en) if 'final_overall_summary_en' in globals() else ""
    document.add_paragraph(
        "The comprehensive insights presented in this report form a robust foundation for further analytical work. "
        "Understanding each variable's behavior—from central tendencies and dispersion to seasonal patterns and anomalies—enables targeted data cleaning, transformation, and predictive modeling. "
        "Advanced visualizations and statistical techniques further enhance interpretability and support data-driven decision making. " + final_summary
    )
    
    return document

def generate_doc_report_id(df):
    document = Document()
    
    # --- Pendahuluan dan Statistik Deskriptif Dasar ---
    document.add_heading("Laporan Eksplorasi Data", level=1)
    document.add_paragraph(
        f"Laporan ini memberikan analisis mendalam terhadap dataset yang terdiri dari {df.shape[0]} baris dan {df.shape[1]} kolom. "
        "Analisis ini mencakup statistik dasar, struktur data, serta analisis univariat dan bivariat. "
        "Wawasan yang diperoleh bisa digunakan untuk pembersihan data, transformasi, dan pemodelan prediktif."
    )
    
    # --- Jenis Data dan Struktur ---
    document.add_heading("Jenis Data dan Struktur", level=2)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    # Misalnya, kolom tanggal tidak digunakan dalam analisis statistik ini
    datetime_cols = []  # Jika ada kolom tanggal yang ingin dikecualikan, masukkan di sini
    numeric_cols = [col for col in numeric_cols if col not in datetime_cols]
    object_cols = df.select_dtypes(include='object').columns.tolist()
    bool_cols = df.select_dtypes(include='bool').columns.tolist()
    category_cols = df.select_dtypes(include='category').columns.tolist()
    document.add_paragraph(
        "Dataset ini memiliki berbagai tipe data. Variabel numerik (misalnya, " +
        (", ".join(numeric_cols) if numeric_cols else "tidak ada") +
        ") memberikan informasi kuantitatif yang penting, sedangkan variabel kategorikal dan boolean memberikan konteks kualitatif. "
        "Analisis ini akan menggunakan statistik dasar untuk menggambarkan karakteristik data secara keseluruhan."
    )
    
    # --- Analisis Univariat dan Deteksi Outlier (kecuali kolom tanggal) ---
    document.add_heading("Analisis Univariat dan Deteksi Outlier", level=2)
    if numeric_cols:
        for col in numeric_cols:
            document.add_heading(f"Analisis Variabel Numerik: '{col}'", level=3)
            series = df[col].dropna()
            if series.empty:
                document.add_paragraph(f"Tidak ada data untuk '{col}' setelah mengeluarkan nilai yang hilang.")
                continue
            mean_val = series.mean()
            median_val = series.median()
            try:
                mode_val = series.mode()[0]
            except Exception:
                mode_val = "Tidak Tersedia"
            std_val = series.std()
            var_val = series.var()
            min_val = series.min()
            max_val = series.max()
            range_val = max_val - min_val
            q25 = series.quantile(0.25)
            q75 = series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = series.skew()
            kurt_val = series.kurt()
            lower_bound = q25 - 1.5 * iqr_val
            upper_bound = q75 + 1.5 * iqr_val
            outlier_count = series[(series < lower_bound) | (series > upper_bound)].count()
            distribution = "seimbang" if abs(mean_val - median_val) < std_val * 0.1 else "condong"
            document.add_paragraph(
                f"Variabel '{col}' memiliki rata-rata {mean_val:.4f}, median {median_val:.4f}, dan modus {mode_val}. "
                f"Nilai standar deviasi adalah {std_val:.4f} dan variansnya {var_val:.4f}, menghasilkan rentang keseluruhan {range_val:.4f}. "
                f"IQR-nya adalah {iqr_val:.4f} dengan Q1 {q25:.4f} dan Q3 {q75:.4f}. "
                f"Skewness tercatat {skew_val:.4f} dan kurtosis {kurt_val:.4f}. Terdeteksi {outlier_count} outlier, yang memberikan gambaran lengkap mengenai sebaran data."
            )
    else:
        document.add_paragraph("Tidak ada variabel numerik yang tersedia untuk analisis univariat.")
    
    # --- Analisis Korelasi Bivariat ---
    document.add_heading("Analisis Korelasi Bivariat", level=2)
    if numeric_cols and len(numeric_cols) > 1:
        corr_matrix = df[numeric_cols].corr()
        for col in numeric_cols:
            document.add_heading(f"Analisis Korelasi untuk '{col}'", level=3)
            corr_series = corr_matrix[col].drop(labels=[col]).sort_values(ascending=False)
            top_n = min(5, len(corr_series))
            if top_n > 0:
                top_corr = corr_series.iloc[:top_n]
                corr_text = ", ".join([f"{idx} ({corr:.4f})" for idx, corr in top_corr.items()])
                document.add_paragraph(
                    f"Untuk '{col}', analisis menunjukkan bahwa {top_n} variabel dengan korelasi positif tertinggi adalah: {corr_text}. "
                    "Korelasi ini mengindikasikan hubungan antar variabel yang bisa berguna untuk pemodelan prediktif."
                )
            else:
                document.add_paragraph(f"Tidak terdeteksi korelasi signifikan untuk '{col}'.")
    else:
        document.add_paragraph("Analisis korelasi bivariat tidak dapat dilakukan karena jumlah variabel numerik yang tersedia tidak mencukupi.")
    
    # --- Wawasan Variabel dan Implikasi Keseluruhan ---
    document.add_heading("Wawasan Variabel dan Implikasi Keseluruhan", level=2)
    overall_intro = random.choice(overall_insight_intro_id) if 'overall_insight_intro_id' in globals() else ""
    document.add_paragraph(
        overall_intro + " Dalam bagian akhir ini, setiap variabel dianalisis secara mendalam untuk mengekstraksi wawasan yang dapat dijadikan dasar analisis lebih lanjut. "
        "Dengan menggabungkan ukuran statistik dan pola yang terlihat, kita mendapatkan pemahaman lengkap tentang struktur, variabilitas, dan anomali data, "
        "yang dapat digunakan untuk pembersihan, transformasi, dan pemodelan prediktif."
    )
    
    for col in df.columns:
        # Pastikan kolom tanggal asli tidak ikut dianalisis
        if col in datetime_cols or col == "date_val":
            continue
        document.add_heading(f"Wawasan Detail untuk '{col}'", level=3)
        var_detail_intro = random.choice(variables_detail_intro_en) if 'variables_detail_intro_en' in globals() else ""
        document.add_paragraph(var_detail_intro)
        col_series = df[col]
        base_info = (
            f"Tipe data: {col_series.dtype}. Variabel ini memiliki {col_series.nunique()} nilai unik dan {col_series.isnull().sum()} entri yang hilang. "
        )
        if pd.api.types.is_numeric_dtype(col_series):
            mean_val = col_series.mean()
            median_val = col_series.median()
            std_val = col_series.std()
            var_val = col_series.var()
            min_val = col_series.min()
            max_val = col_series.max()
            range_val = max_val - min_val
            q25 = col_series.quantile(0.25)
            q75 = col_series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = col_series.skew()
            kurt_val = col_series.kurt()
            distribution = "seimbang" if abs(mean_val - median_val) < std_val * 0.1 else "condong"
            descriptive_text = random.choice(descriptive_stats_template_id).format(
                mean_val=mean_val, median_val=median_val, mode_val="N/A", distribution=distribution
            ) if 'descriptive_stats_template_id' in globals() else f"Nilai rata-ratanya adalah {mean_val:.4f} dan mediannya {median_val:.4f}."
            variability_text = random.choice(variability_insight_id).format(std_val=std_val, var_val=var_val) if 'variability_insight_id' in globals() else ""
            detailed_info = (
                f"Untuk variabel numerik ini, {descriptive_text} Standar deviasinya adalah {std_val:.4f} dan variansnya {var_val:.4f}, yang menunjukkan dispersi yang signifikan. {variability_text}"
            )
        elif pd.api.types.is_object_dtype(col_series) or pd.api.types.is_bool_dtype(col_series):
            try:
                mode_val = col_series.mode()[0]
                detailed_info = (
                    f"Untuk variabel kategorikal ini, nilai modus yang paling sering muncul adalah '{mode_val}', yang menunjukkan adanya kategori dominan yang bisa mempengaruhi analisis."
                )
            except Exception:
                detailed_info = "Modus tidak dapat ditentukan untuk variabel ini."
        else:
            detailed_info = "Analisis statistik mendalam tidak dapat dilakukan untuk tipe variabel ini."
        document.add_paragraph(base_info + detailed_info)
    
    final_summary = random.choice(final_overall_summary_id) if 'final_overall_summary_id' in globals() else ""
    document.add_paragraph(
        "Wawasan komprehensif yang dipaparkan dalam laporan ini menjadi dasar yang kuat untuk analisis lebih lanjut. "
        "Memahami perilaku tiap variabel—dari tendensi sentral dan dispersi hingga pola dan anomali—memungkinkan pembersihan data, transformasi, "
        "dan pemodelan prediktif yang lebih terarah. Visualisasi dan teknik statistik lanjutan semakin meningkatkan interpretabilitas dan mendukung pengambilan keputusan berbasis data. " + final_summary
    )
    
    return document

def generate_doc_report_id_ts(df):
    # Duplikasi kolom datetime terpilih ke dalam kolom baru "date_val"
    if ts_choice in df.columns:
        df["date_val"] = df[ts_choice]
    
    # Konversi kolom datetime yang terpilih menjadi objek datetime dan jadikan sebagai index
    if ts_choice in df.columns:
        df[ts_choice] = pd.to_datetime(df[ts_choice], errors='raise')
        df = df.set_index(ts_choice)
    
    document = Document()
    
    # --- Pendahuluan dan Statistik Deskriptif Dasar ---
    document.add_heading("Laporan Eksplorasi Data", level=1)
    document.add_paragraph(
        f"Laporan ini menyajikan analisis mendalam terhadap dataset yang memiliki {df.shape[0]} baris dan {df.shape[1]} kolom. "
        "Analisis deret waktu ini dirancang untuk menunjukkan tren jangka panjang, pola siklik yang halus, dan perubahan mendadak selama periode pengamatan yang luas. "
        "Informasi ini sangat berguna untuk peramalan dan pengambilan keputusan yang tepat."
    )
    
    # --- Jenis Data dan Struktur ---
    document.add_heading("Jenis Data dan Struktur", level=2)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    # Kecualikan kolom index datetime dan kolom 'date_val'
    datetime_cols = [ts_choice, "date_val"]
    numeric_cols = [col for col in numeric_cols if col not in datetime_cols]
    object_cols = df.select_dtypes(include='object').columns.tolist()
    bool_cols = df.select_dtypes(include='bool').columns.tolist()
    category_cols = df.select_dtypes(include='category').columns.tolist()
    document.add_paragraph(
        "Dataset ini memiliki beragam tipe data. Variabel numerik (misalnya, " +
        (", ".join(numeric_cols) if numeric_cols else "tidak ada") +
        ") memberikan informasi kuantitatif yang penting, sedangkan variabel kategorikal dan boolean memberikan konteks kualitatif. "
        f"Kolom datetime '{ts_choice}' telah dikonversi ke format yang benar dan dijadikan sebagai index, dan nilai aslinya disimpan di kolom 'date_val' sebagai referensi."
    )
    
    # --- Analisis Univariat dan Deteksi Outlier (kecuali kolom datetime dan date_val) ---
    document.add_heading("Analisis Univariat dan Deteksi Outlier", level=2)
    if numeric_cols:
        for col in numeric_cols:
            document.add_heading(f"Analisis Variabel Numerik: '{col}'", level=3)
            series = df[col].dropna()
            if series.empty:
                document.add_paragraph(f"Tidak ada data untuk '{col}' setelah mengeluarkan nilai yang hilang.")
                continue
            mean_val = series.mean()
            median_val = series.median()
            try:
                mode_val = series.mode()[0]
            except Exception:
                mode_val = "Tidak Tersedia"
            std_val = series.std()
            var_val = series.var()
            min_val = series.min()
            max_val = series.max()
            range_val = max_val - min_val
            q25 = series.quantile(0.25)
            q75 = series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = series.skew()
            kurt_val = series.kurt()
            lower_bound = q25 - 1.5 * iqr_val
            upper_bound = q75 + 1.5 * iqr_val
            outlier_count = series[(series < lower_bound) | (series > upper_bound)].count()
            distribution = "seimbang" if abs(mean_val - median_val) < std_val * 0.1 else "condong"
            document.add_paragraph(
                f"Variabel '{col}' memiliki rata-rata {mean_val:.4f}, median {median_val:.4f}, dan modus {mode_val}. "
                f"Standar deviasi tercatat {std_val:.4f} dan varians {var_val:.4f}, sehingga rentang keseluruhan adalah {range_val:.4f}. "
                f"IQR-nya adalah {iqr_val:.4f} dengan Q1 {q25:.4f} dan Q3 {q75:.4f}. "
                f"Skewness adalah {skew_val:.4f} dan kurtosis {kurt_val:.4f}. Terdeteksi {outlier_count} outlier, yang memberikan gambaran lengkap mengenai sebaran data."
            )
    else:
        document.add_paragraph("Tidak ada variabel numerik yang tersedia untuk analisis univariat.")
    
    # --- Analisis Korelasi Bivariat ---
    document.add_heading("Analisis Korelasi Bivariat", level=2)
    if numeric_cols and len(numeric_cols) > 1:
        corr_matrix = df[numeric_cols].corr()
        for col in numeric_cols:
            document.add_heading(f"Analisis Korelasi untuk '{col}'", level=3)
            corr_series = corr_matrix[col].drop(labels=[col]).sort_values(ascending=False)
            top_n = min(5, len(corr_series))
            if top_n > 0:
                top_corr = corr_series.iloc[:top_n]
                corr_text = ", ".join([f"{idx} ({corr:.4f})" for idx, corr in top_corr.items()])
                document.add_paragraph(
                    f"Untuk '{col}', terdapat {top_n} variabel dengan korelasi positif tertinggi: {corr_text}. "
                    "Korelasi ini menunjukkan adanya hubungan antar variabel yang bisa berguna untuk pemodelan prediktif."
                )
            else:
                document.add_paragraph(f"Tidak ada korelasi signifikan untuk '{col}'.")
    else:
        document.add_paragraph("Analisis korelasi bivariat tidak dapat dilakukan karena jumlah variabel numerik yang tersedia tidak cukup.")
    
    # --- Analisis Deret Waktu ---
    document.add_heading("Analisis Deret Waktu", level=2)
    # Gunakan kolom 'date_val' untuk mendapatkan nilai asli tanggal
    if "date_val" in df.columns:
        start_date = df["date_val"].iloc[0]
        end_date = df["date_val"].iloc[-1]
    else:
        start_date = df.index[0]
        end_date = df.index[-1]
    ts_intro = random.choice(time_series_analysis_intro_en) if 'time_series_analysis_intro_en' in globals() else (
        "Deret waktu mencakup periode dari {start_date} hingga {end_date}, menyediakan kerangka waktu yang kuat untuk analisis."
    )
    document.add_paragraph(ts_intro.format(start_date=start_date, end_date=end_date))
    
    # --- Analisis Pola Musiman menggunakan Frekuensi yang Ditetapkan Pengguna ---
    document.add_heading("Analisis Pola Musiman", level=3)
    freq_mapping = {"Seconds": "S", "Minutes": "T", "Hours": "H", "Days": "D", "Weeks": "W", "Months": "M", "Quarterly": "Q", "Yearly": "Y"}
    freq_code = freq_mapping.get(timefreq, "D")
    season_text_full = (
        f"Dengan melakukan resampling data menggunakan frekuensi terpilih ({timefreq}), kita dapat melihat pola musiman tiap variabel numerik. "
        "Analisis ini menghitung rata-rata keseluruhan, standar deviasi, serta kuartil (Q1, median, Q3) untuk setiap periode. "
        "Jika nilai-nilai kuartil mendekati rata-rata, berarti pola musiman cenderung stabil, namun jika ada jarak yang lebar, itu mengindikasikan variabilitas yang tinggi. "
        "Analisis ini memberikan gambaran mendalam tentang bagaimana data berperilaku tiap periode. \n\n"
    )
    for col in numeric_cols:
        try:
            resampled = df[col].resample(freq_code).mean()
            if resampled.count() < 2:
                season_text_full += f"Untuk '{col}', data dalam setiap {timefreq.lower()} terlalu sedikit untuk menghasilkan statistik yang bermakna.\n\n"
                continue
            overall_avg = resampled.mean()
            overall_std = resampled.std()
            q1 = resampled.quantile(0.25)
            median = resampled.quantile(0.50)
            q3 = resampled.quantile(0.75)
            season_text_full += (
                f"Untuk '{col}':\n"
                f"Dari data yang telah diresampling, rata-rata keseluruhan adalah {overall_avg:.4f} dengan standar deviasi {overall_std:.4f}. "
                f"Nilai kuartil pertama (Q1) adalah {q1:.4f}, median {median:.4f}, dan kuartil ketiga (Q3) adalah {q3:.4f}. "
                "Jika nilai-nilai tersebut terkonsentrasi, pola musiman cenderung stabil; jika tersebar, terdapat fluktuasi musiman yang cukup besar. \n\n"
            )
        except Exception as e:
            season_text_full += f"Analisis musiman untuk '{col}' tidak dapat diselesaikan karena error: {e}.\n\n"
    document.add_paragraph(season_text_full)
    
    # --- Analisis Statistik Bergulir dan Tren ---
    document.add_heading("Analisis Statistik Bergulir dan Tren", level=3)
    roll_text_full = ""
    for col in numeric_cols:
        try:
            window_results = []
            for window in range(1, 13):
                roll_mean_val = df[col].rolling(window=window, min_periods=1).mean().mean()
                roll_std_val = df[col].rolling(window=window, min_periods=1).std().mean()
                window_results.append((window, roll_mean_val, roll_std_val))
            best_mean = max(window_results, key=lambda x: x[1])
            worst_mean = min(window_results, key=lambda x: x[1])
            highest_std = max([r for r in window_results if not np.isnan(r[2])], key=lambda x: x[2], default=(None, None, None))
            lowest_std = min([r for r in window_results if not np.isnan(r[2])], key=lambda x: x[2], default=(None, None, None))
            roll_text_full += (
                f"Untuk '{col}':\n"
                f"Dengan mencoba ukuran jendela dari 1 sampai 12 {timefreq.lower()}(s), diperoleh bahwa rata-rata bergulir tertinggi sebesar {best_mean[1]:.4f} terjadi dengan jendela berukuran {best_mean[0]} {timefreq.lower()}(s), sedangkan rata-rata bergulir terendah sebesar {worst_mean[1]:.4f} terjadi dengan jendela berukuran {worst_mean[0]} {timefreq.lower()}(s). "
                f"Dari segi volatilitas, standar deviasi bergulir tertinggi adalah {highest_std[2]:.4f} pada jendela {highest_std[0]} {timefreq.lower()}(s), dan yang terendah adalah {lowest_std[2]:.4f} pada jendela {lowest_std[0]} {timefreq.lower()}(s). "
                "Analisis ini menunjukkan bagaimana perilaku variabel berubah saat dilihat dari tingkat agregasi yang berbeda, menyoroti fluktuasi jangka pendek dan tren jangka panjang. \n\n"
            )
        except Exception as e:
            roll_text_full += f"Analisis statistik bergulir untuk '{col}' tidak dapat dihitung karena error: {e}.\n\n"
    document.add_paragraph(roll_text_full)
    
    # --- Analisis Pola Deret Waktu Lanjutan ---
    document.add_heading("Analisis Pola Deret Waktu Lanjutan", level=3)
    adv_text_full = ""
    for col in numeric_cols:
        try:
            if df[col].dropna().empty:
                adv_text_full += f"Tidak ada data untuk analisis lanjutan variabel '{col}'.\n\n"
                continue
            max_idx = df[col].idxmax()
            min_idx = df[col].idxmin()
            if "date_val" in df.columns:
                max_date = df["date_val"].iloc[max_idx] if isinstance(max_idx, int) and max_idx < len(df) else max_idx
                min_date = df["date_val"].iloc[min_idx] if isinstance(min_idx, int) and min_idx < len(df) else min_idx
            else:
                max_date = max_idx
                min_date = min_idx
            max_val = df[col].max()
            min_val = df[col].min()
            ts_detail = (
                f"Untuk '{col}', nilai tertinggi yang diamati adalah {max_val:.4f} yang tercatat pada {max_date}, dan nilai terendah yang diamati adalah {min_val:.4f} yang tercatat pada {min_date}. "
            )
            try:
                adf_result = adfuller(df[col].dropna())
                p_value = adf_result[1]
                stationarity = "stasioner" if p_value < 0.05 else "tidak stasioner"
                ts_detail += f"Uji Augmented Dickey-Fuller (ADF) menunjukkan bahwa deret ini {stationarity} dengan p-value sebesar {p_value:.4f}."
            except Exception as adf_err:
                ts_detail += f"Uji ADF tidak dapat dilakukan: {adf_err}."
            adv_text_full += ts_detail + "\n\n"
        except Exception as e:
            adv_text_full += f"Analisis pola deret waktu lanjutan untuk '{col}' mengalami error: {e}.\n\n"
    document.add_paragraph(adv_text_full)
    
    # --- Wawasan Variabel dan Implikasi Keseluruhan ---
    document.add_heading("Wawasan Variabel dan Implikasi Keseluruhan", level=2)
    overall_intro = random.choice(overall_insight_intro_en) if 'overall_insight_intro_en' in globals() else ""
    document.add_paragraph(
        overall_intro + " Dalam bagian akhir ini, setiap variabel dianalisis secara mendalam untuk mengekstraksi wawasan yang dapat menjadi dasar analisis lebih lanjut. "
        "Dengan menggabungkan ukuran statistik kunci dan pola yang teramati, kita mendapatkan pemahaman menyeluruh tentang struktur, variabilitas, dan anomali dataset, "
        "yang dapat digunakan untuk pembersihan data, transformasi, dan pemodelan prediktif."
    )
    
    for col in df.columns:
        if col == ts_choice or col == "date_val":
            continue
        document.add_heading(f"Wawasan Detail untuk '{col}'", level=3)
        var_detail_intro = random.choice(variables_detail_intro_en) if 'variables_detail_intro_en' in globals() else ""
        document.add_paragraph(var_detail_intro)
        col_series = df[col]
        base_info = (
            f"Tipe data: {col_series.dtype}. Variabel ini memiliki {col_series.nunique()} nilai unik dan {col_series.isnull().sum()} entri yang hilang. "
        )
        if pd.api.types.is_numeric_dtype(col_series):
            mean_val = col_series.mean()
            median_val = col_series.median()
            std_val = col_series.std()
            var_val = col_series.var()
            min_val = col_series.min()
            max_val = col_series.max()
            range_val = max_val - min_val
            q25 = col_series.quantile(0.25)
            q75 = col_series.quantile(0.75)
            iqr_val = q75 - q25
            skew_val = col_series.skew()
            kurt_val = col_series.kurt()
            distribution = "seimbang" if abs(mean_val - median_val) < std_val * 0.1 else "condong"
            descriptive_text = random.choice(descriptive_stats_template_en).format(
                mean_val=mean_val, median_val=median_val, mode_val="N/A", distribution=distribution
            ) if 'descriptive_stats_template_en' in globals() else f"Nilai rata-ratanya adalah {mean_val:.4f} dan mediannya {median_val:.4f}."
            variability_text = random.choice(variability_insight_en).format(std_val=std_val, var_val=var_val) if 'variability_insight_en' in globals() else ""
            detailed_info = (
                f"Untuk variabel numerik ini, {descriptive_text} Standar deviasinya adalah {std_val:.4f} dan variansnya {var_val:.4f}, yang menunjukkan dispersi yang signifikan. {variability_text}"
            )
        elif pd.api.types.is_object_dtype(col_series) or pd.api.types.is_bool_dtype(col_series):
            try:
                mode_val = col_series.mode()[0]
                detailed_info = (
                    f"Untuk variabel kategorikal ini, nilai modus yang paling sering muncul adalah '{mode_val}', yang menunjukkan adanya kategori dominan yang dapat mempengaruhi analisis."
                )
            except Exception:
                detailed_info = "Modus tidak dapat ditentukan untuk variabel ini."
        else:
            detailed_info = "Analisis statistik mendalam tidak dapat dilakukan untuk tipe variabel ini."
        document.add_paragraph(base_info + detailed_info)
    
    final_summary = random.choice(final_overall_summary_en) if 'final_overall_summary_en' in globals() else ""
    document.add_paragraph(
        "Wawasan komprehensif yang dipaparkan dalam laporan ini membentuk dasar yang kuat untuk analisis lebih lanjut. "
        "Memahami perilaku tiap variabel—dari tendensi sentral dan dispersi hingga pola dan anomali—memungkinkan pembersihan data, transformasi, "
        "dan pemodelan prediktif yang lebih terarah. Visualisasi dan teknik statistik lanjutan semakin meningkatkan interpretabilitas dan mendukung pengambilan keputusan berbasis data. " + final_summary
    )
    
    return document

# ============================================== Download Functions =========================================================================================

def aggregate_and_download_plots(df, selected_plot_types, theme="Blue", cmap_option="Blue"):
    """Aggregates figures from the selected plot types and returns a zip buffer."""
    all_figures = []

    # Normalize plot type names for comparison
    selected_plot_types_lower = [plot_type.lower() for plot_type in selected_plot_types]

    if "histograms" in selected_plot_types_lower:
        figs = get_histogram_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "boxplots" in selected_plot_types_lower:
        figs = get_boxplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "scatterplots" in selected_plot_types_lower:
        figs = get_scatterplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "lineplots" in selected_plot_types_lower:
        figs = get_lineplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "areaplots" in selected_plot_types_lower:
        figs = get_areaplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "violinplots" in selected_plot_types_lower:
        figs = get_violinplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "correlation heatmap" in selected_plot_types_lower:
        figs = get_correlation_heatmap_figure(df, cmap_option)
        if figs:
            all_figures.extend(figs)
    if "cdf" in selected_plot_types_lower:
        figs = get_cdf_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "categorical barplots" in selected_plot_types_lower:
        figs = get_categorical_barplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "piecharts" in selected_plot_types_lower:
        figs = get_piechart_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "stacked barplots" in selected_plot_types_lower:
        figs = get_stacked_barplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "grouped barplots" in selected_plot_types_lower:
        figs = get_grouped_barplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "wordclouds" in selected_plot_types_lower:
        figs = get_wordcloud_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "countplots" in selected_plot_types_lower:
        figs = get_countplot_figures(df, theme)
        if figs:
            all_figures.extend(figs)
    if "treemaps" in selected_plot_types_lower:
        figs = get_treemap_figures(df, theme)
        if figs:
            all_figures.extend(figs)

    if not all_figures:
        st.write("No plot types selected or no figures were generated.")
        return None  # Return None if no figures

    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for filename, fig in all_figures:
            buf = io.BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight")
            buf.seek(0)
            zip_file.writestr(filename, buf.getvalue())
            plt.close(fig)  # Close figure to avoid memory warning.

    zip_buffer.seek(0)
    return zip_buffer  # return the zip buffer.

# Plotting for Report Data ================================================================================================================================================================================

def eda_dataframe_to_docx_en(df):
    """
    Performs Exploratory Data Analysis (EDA) on a DataFrame and creates a docx document.
    """
    document = Document()

    document.add_heading("Exploratory Data Analysis", 0)

    document.add_heading("Basic Statistics", level=1)
    document.add_paragraph(f"Number of Observations (Rows): {df.shape[0]}")
    document.add_paragraph(f"Number of Variables (Columns): {df.shape[1]}")
    document.add_paragraph(f"Duplicate Rows: {df.duplicated().sum()}")
    document.add_paragraph(f"Missing Cells: {df.isnull().sum().sum()}")

    document.add_heading("Variable Types", level=1)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    categorical_cols = df.select_dtypes(include='category').columns.tolist()
    text_cols = df.select_dtypes(include='object').columns.tolist()

    document.add_paragraph(f"Numeric: {len(numeric_cols)} ({', '.join(map(str, numeric_cols))})")
    document.add_paragraph(f"Categorical: {len(categorical_cols)} ({', '.join(map(str, categorical_cols))})")
    document.add_paragraph(f"Text (Object): {len(text_cols)} ({', '.join(map(str, text_cols))})")

    document.add_heading("Variables with Unique Values", level=1)
    unique_counts = df.nunique()
    document.add_paragraph(str(unique_counts))

    document.add_heading("Variables with Uniform Distribution (Simplified)", level=1)
    for col in df.columns:
        if df[col].nunique() > 10:
            if (df[col].value_counts(normalize=True).std() < 0.05):
                document.add_paragraph(f"Column: {col} might have a uniform distribution")
        else:
            document.add_paragraph(f"Column: {col} has less than 10 unique values, cannot check for uniform distribution")

    document.add_heading("Missing Values per Variable", level=1)
    document.add_paragraph(str(df.isnull().sum()))

    document.add_heading("Top 5 Mostly and Least Correlated Variables", level=1)
    if len(numeric_cols) > 0:
        corr_matrix = df[numeric_cols].corr().abs()
        for col in numeric_cols:
            if col in corr_matrix.columns:
                corr_series = corr_matrix[col].sort_values(ascending=False)
                document.add_paragraph(f"\nCorrelation with {col}:")
                document.add_paragraph("Top 5 Mostly Correlated:")
                document.add_paragraph(str(corr_series.head(6).tail(5)))
                document.add_paragraph("Top 5 Least Correlated:")
                document.add_paragraph(str(corr_series.tail(5)))

    return document

def eda_dataframe_to_docx_id(df):
    """
    Melakukan Analisis Data Eksplorasi (EDA) pada DataFrame dan membuat dokumen docx dalam bahasa Indonesia.
    """
    document = Document()

    document.add_heading("Analisis Data Eksplorasi", 0)

    document.add_heading("Statistik Dasar", level=1)
    document.add_paragraph(f"Jumlah Observasi (Baris): {df.shape[0]}")
    document.add_paragraph(f"Jumlah Variabel (Kolom): {df.shape[1]}")
    document.add_paragraph(f"Baris Duplikat: {df.duplicated().sum()}")
    document.add_paragraph(f"Sel yang Hilang: {df.isnull().sum().sum()}")

    document.add_heading("Tipe Variabel", level=1)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    categorical_cols = df.select_dtypes(include='category').columns.tolist()
    text_cols = df.select_dtypes(include='object').columns.tolist()

    document.add_paragraph(f"Numerik: {len(numeric_cols)} ({', '.join(map(str, numeric_cols))})")
    document.add_paragraph(f"Kategorikal: {len(categorical_cols)} ({', '.join(map(str, categorical_cols))})")
    document.add_paragraph(f"Teks (Objek): {len(text_cols)} ({', '.join(map(str, text_cols))})")

    document.add_heading("Variabel dengan Nilai Unik", level=1)
    unique_counts = df.nunique()
    document.add_paragraph(str(unique_counts))

    document.add_heading("Variabel dengan Distribusi Seragam (Sederhana)", level=1)
    for col in df.columns:
        if df[col].nunique() > 10:
            if (df[col].value_counts(normalize=True).std() < 0.05):
                document.add_paragraph(f"Kolom: {col} mungkin memiliki distribusi seragam")
        else:
            document.add_paragraph(f"Kolom: {col} memiliki kurang dari 10 nilai unik, tidak dapat diperiksa distribusi seragam")

    document.add_heading("Nilai yang Hilang per Variabel", level=1)
    document.add_paragraph(str(df.isnull().sum()))

    document.add_heading("5 Variabel dengan Korelasi Tertinggi dan Terendah", level=1)
    if len(numeric_cols) > 0:
        corr_matrix = df[numeric_cols].corr().abs()
        for col in numeric_cols:
            if col in corr_matrix.columns:
                corr_series = corr_matrix[col].sort_values(ascending=False)
                document.add_paragraph(f"\nKorelasi dengan {col}:")
                document.add_paragraph("5 Korelasi Tertinggi:")
                document.add_paragraph(str(corr_series.head(6).tail(5)))
                document.add_paragraph("5 Korelasi Terendah:")
                document.add_paragraph(str(corr_series.tail(5)))

    return document

def eda_all_columns_to_docx_en(df):
    """
    Performs Exploratory Data Analysis on all columns iteratively and creates a docx document.
    """
    document = Document()
    document.add_heading("Column-Wise Exploratory Data Analysis", 0)

    for column_name in df.columns:
        col = df[column_name]
        document.add_heading(f"Analysis of Column: {column_name}", level=1)
        document.add_paragraph(f"Data Type: {col.dtype}")
        document.add_paragraph(f"Number of Unique Values: {col.nunique()}")
        document.add_paragraph(f"Number of Missing Values: {col.isnull().sum()}")

        if pd.api.types.is_numeric_dtype(col):
            document.add_paragraph(f"Mean: {col.mean()}")
            document.add_paragraph(f"Standard Deviation: {col.std()}")
            document.add_paragraph(f"Minimum: {col.min()}")
            document.add_paragraph(f"25th Percentile: {col.quantile(0.25)}")
            document.add_paragraph(f"Median: {col.median()}")
            document.add_paragraph(f"75th Percentile: {col.quantile(0.75)}")
            document.add_paragraph(f"Maximum: {col.max()}")
            document.add_paragraph(f"Skewness: {col.skew()}")
            document.add_paragraph(f"Kurtosis: {col.kurt()}")
            document.add_paragraph(f"Number of Zeros: {(col == 0).sum()}")

        elif pd.api.types.is_string_dtype(col) or pd.api.types.is_object_dtype(col):
            most_frequent = col.mode()[0]
            document.add_paragraph(f"Most Frequent Value: {most_frequent}")
            document.add_paragraph(f"Frequency of Most Frequent Value: {(col == most_frequent).sum()}")

    return document

def eda_all_columns_to_docx_id(df):
    """
    Melakukan Analisis Data Eksplorasi pada semua kolom secara iteratif dan membuat dokumen docx.
    """
    document = Document()
    document.add_heading("Analisis Data Eksplorasi Per Kolom", 0)

    for column_name in df.columns:
        col = df[column_name]
        document.add_heading(f"Analisis Kolom: {column_name}", level=1)
        document.add_paragraph(f"Tipe Data: {col.dtype}")
        document.add_paragraph(f"Jumlah Nilai Unik: {col.nunique()}")
        document.add_paragraph(f"Jumlah Nilai Hilang: {col.isnull().sum()}")

        if pd.api.types.is_numeric_dtype(col):
            document.add_paragraph(f"Rata-rata: {col.mean()}")
            document.add_paragraph(f"Standar Deviasi: {col.std()}")
            document.add_paragraph(f"Minimum: {col.min()}")
            document.add_paragraph(f"Persentil ke-25: {col.quantile(0.25)}")
            document.add_paragraph(f"Median: {col.median()}")
            document.add_paragraph(f"Persentil ke-75: {col.quantile(0.75)}")
            document.add_paragraph(f"Maksimum: {col.max()}")
            document.add_paragraph(f"Kemiringan (Skewness): {col.skew()}")
            document.add_paragraph(f"Kurtosis: {col.kurt()}")
            document.add_paragraph(f"Jumlah Nol: {(col == 0).sum()}")

        elif pd.api.types.is_string_dtype(col) or pd.api.types.is_object_dtype(col):
            most_frequent = col.mode()[0]
            document.add_paragraph(f"Nilai Paling Sering Muncul: {most_frequent}")
            document.add_paragraph(f"Frekuensi Nilai Paling Sering Muncul: {(col == most_frequent).sum()}")

    return document

# ============================================ QUOTA HANDLING =============================================================================

def get_login_df():
    """Retrieves login data from the 'Accounts' sheet."""
    url = f'https://docs.google.com/spreadsheet/ccc?key={sheetskey}&output=csv'
    df_login = pd.read_csv(url)
    return df_login

def send_email(email, message):
    """Sends an email notification."""
    try:
        sender_email = "insightfoxa@gmail.com"  # Replace with your sender email
        sender_password = "txsr udsn fhpe pfns"  # Replace with your app password

        msg = MIMEText(message)
        msg['Subject'] = 'Password Reminder'
        msg['From'] = sender_email
        msg['To'] = email

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, email, msg.as_string())

        st.success("Reminder email sent. Please check your inbox (and spam folder)!")
    except Exception as e:
        st.error(f"Failed to send email: {e}")

def get_quota_df():
    """Retrieves quota data from the 'Quotas' sheet."""
    sheet = client.open_by_key(sheetskey).worksheet("Quotas")
    data = sheet.get_all_records()
    return pd.DataFrame(data)

def update_quota(email, column):
    """
    Decrements the quota in the given column for the specified email.
    If quota is >0, subtracts 1 and updates the Google Sheet.
    """
    sheet = client.open_by_key(sheetskey).worksheet("Quotas")
    df = get_quota_df()
    user_idx = df.index[df['email'] == email].tolist()
    if user_idx:
        row_index = user_idx[0] + 2  # Adjust for header row and 1-indexing in Sheets
        current_quota = int(df.loc[user_idx[0], column])
        if current_quota > 0:
            new_quota = current_quota - 1
            # Compute the column letter (assumes columns are in order)
            col_idx = df.columns.get_loc(column)  # 0-based index
            col_letter = chr(65 + col_idx)  # 0 -> A, 1 -> B, etc.
            sheet.update_acell(f"{col_letter}{row_index}", new_quota)
        else:
            st.error("No quota remaining for this action.")
    else:
        st.error("User not found in quota records.")

# ========================================== PAYMENT PROCESSING =====================================================================================

def colnum_to_letters(n):
    """Convert a 1-indexed column number to Excel-style column letters."""
    result = ""
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        result = chr(65 + remainder) + result
    return result

def get_gemini_repsonse(input_prompt, image_parts):
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([input_prompt, image_parts[0]])
    return response.text

def input_image_setup(uploaded_file):
    # Read the uploaded file into bytes
    bytes_data = uploaded_file.read()
    image_parts = [
        {
            "mime_type": uploaded_file.type,
            "data": bytes_data
        }
    ]
    return image_parts

def extract_amount(text):
    """Extracts currency and amount from the Gemini response."""
    try:
        pattern = r"([A-Z]{3})\s*-\s*([\d.]+)"
        matches = re.findall(pattern, text)
        if matches:
            return matches
        else:
            return None
    except Exception as e:
        st.error(f"Error extracting amount: {e}")
        return None

def get_quota_df():
    """Retrieves quota data from the 'Quotas' sheet."""
    sheet = client.open_by_key(sheetskey).worksheet("Quotas")
    data = sheet.get_all_records()
    return pd.DataFrame(data)

def add_quota(email, column, amount_to_add=1):
    """
    Updates the quota in the given column for the specified email by adding the specified amount.
    """
    sheet = client.open_by_key(sheetskey).worksheet("Quotas")
    df = get_quota_df()
    user_idx = df.index[df['email'] == email].tolist()
    if user_idx:
        row_index = user_idx[0] + 2  # Adjust for header row and 1-indexing in Sheets
        current_quota = int(df.loc[user_idx[0], column])
        new_quota = current_quota + amount_to_add
        col_idx = df.columns.get_loc(column) + 1  # convert 0-indexed to 1-indexed
        col_letter = colnum_to_letters(col_idx)
        sheet.update_acell(f"{col_letter}{row_index}", new_quota)
    else:
        st.error("User not found in quota records.")

def generate_image_hash(time_str):
    """Generates a hash of the transaction time string."""
    return hashlib.md5(time_str.encode()).hexdigest()

def update_image_hashes(email, image_hash):
    """Updates the image hashes for the specified email in Google Sheets."""
    sheet = client.open_by_key(sheetskey).worksheet("Quotas")
    df = get_quota_df()
    user_idx = df.index[df['email'] == email].tolist()
    if user_idx:
        row_index = user_idx[0] + 2
        # Get the counter cell from the column named 'counter'
        if "counter" in df.columns:
            counter_col_idx = df.columns.get_loc("counter") + 1  # convert 0-indexed to 1-indexed
            counter_col_letter = colnum_to_letters(counter_col_idx)
            counter_cell = sheet.acell(f"{counter_col_letter}{row_index}").value
            counter = int(counter_cell) if counter_cell and counter_cell.isdigit() else 1
        else:
            # If no counter column exists, default to 1
            counter = 1

        # Determine the hash column using the counter.
        hash_column_index = 5 + counter  # 5 is the base index for the first hash column (1-indexed)
        col_letter = colnum_to_letters(hash_column_index)
        sheet.update_acell(f"{col_letter}{row_index}", image_hash)

        counter += 1
        if counter > 25:
            counter = 1
        if "counter" in df.columns:
            sheet.update_acell(f"{counter_col_letter}{row_index}", counter)
    else:
        st.error("User not found in quota records.")

def check_image_hash_exists(email, image_hash):
    """Checks if the image hash already exists for the specified email."""
    sheet = client.open_by_key(sheetskey).worksheet("Quotas")
    df = get_quota_df()
    user_idx = df.index[df['email'] == email].tolist()
    if user_idx:
        row_index = user_idx[0] + 2
        for i in range(1, 26):
            hash_column_index = 5 + i  # 5 is the base index for the first hash column
            col_letter = colnum_to_letters(hash_column_index)
            cell_value = sheet.acell(f"{col_letter}{row_index}").value
            if cell_value == image_hash:
                return True
        return False
    else:
        st.error("User not found in quota records.")
        return False
    
# ========================================== STREAMLIT APP LAYOUT =========================================================================

# ========================================== REGISTRATION PAGE ================================================================================

if st.session_state.page == 1:
    col1_1, col1_2, col1_3 = st.columns([1,3,1])
    with col1_2 : 
        # st header coloring
        st.markdown("<h1 style='text-align: center; color: white;'>InsightFox</h1>", unsafe_allow_html=True)
        st.markdown("<h6 style='text-align: center; color: white;'>Automated One-step Data Reports</h6>", unsafe_allow_html=True)
        with st.expander("", expanded=True):
            reg_email = st.text_input("Email", key="reg_email")
            reg_password = st.text_input("Password", type="password", key="reg_password")
            st.warning("Please use unique passwords and do not register the same email twice.")

            col_reg1, col_reg2 = st.columns([4, 2])
            with col_reg2:
                remind_me = st.checkbox("Remind me via email")
                if st.button("Register"):
                    st.warning("Please wait...")
                    sheet_accounts = client.open_by_key(sheetskey).worksheet("Accounts")
                    sheet_quota = client.open_by_key(sheetskey).worksheet("Quotas")
                    existing_emails = sheet_accounts.col_values(1)
                    if reg_email in existing_emails:
                        st.error("Email already registered.")
                    else:
                        # Append to Accounts sheet
                        sheet_accounts.append_row([reg_email, reg_password])
                        # Append to Quotas sheet with initial quotas (adjust initial values as needed)
                        sheet_quota.append_row([reg_email, 0, 0, 0, 0])
                        st.success("Registration successful!")
                        if remind_me:
                            send_email(reg_email, f"Your account password is: {reg_password}\nPlease keep this information secure.")
            with col_reg1:
                st.write("")
                st.write("")
                st.write("")
                st.write("Already have an account?")
                if st.button("Go to Login Page"):
                    st.session_state.page = 0
                    st.rerun()

# ============================================== LOGIN FUNCTION =================================================================================

elif st.session_state.page == 0 and st.session_state.login == 0:
    col2_1, col2_2, col2_3 = st.columns([1,3,1])
    with col2_2 : 
        # st header coloring
        st.markdown("<h1 style='text-align: center; color: white;'>InsightFox</h1>", unsafe_allow_html=True)
        st.markdown("<h6 style='text-align: center; color: white;'>Automated One-step Data Reports</h6>", unsafe_allow_html=True)
        with st.expander("", expanded=True):
            login_email = st.text_input("Email", key="login_email")
            login_password = st.text_input("Password", type="password", key="login_password")

            col_log1, col_log2 = st.columns([10, 2])
            with col_log2:
                if st.button("Login"):
                    df_login = get_login_df()
                    user_row = df_login[(df_login['email'] == login_email) & (df_login['password'] == login_password)]
                    if not user_row.empty:
                        st.success("Login successful!")
                        st.session_state.login = 1
                        st.session_state.email = login_email
                        st.session_state.page = 2
                        st.rerun()
                    else:
                        st.error("Invalid email or password.")
            with col_log1:
                st.write("")
                st.write("")
                st.write("")
                st.write("Don't have an account? Register")
                if st.button("Go to Register Page"):
                    st.session_state.page = 1
                    st.rerun()

# ================================================= CENTRALIZED NAVBAR =============================================================================

if st.session_state.login == 1:
    colcen1, colcen2 = st.columns([2,7])
    with colcen2 :
        # Centralized navigation widget
        if "selected_page" not in st.session_state:
            st.session_state.selected_page = "Check and Buy Quota"

        pages = ["Data Cleaning/Preprocessing", "Exploratory Data Analysis Report", "Check and Buy Quota"]
        selected = st.segmented_control(
            "",
            pages,
            selection_mode="single",
            key="selected_page"
        )
        # Update the page state based on the selection
        if selected == "Data Cleaning/Preprocessing":
            st.session_state.page = 2
        elif selected == "Exploratory Data Analysis Report":
            st.session_state.page = 3
        elif selected == "Check and Buy Quota":
            st.session_state.page = 4

# ================================================= DATA CLEANING ==========================================================================================

if st.session_state.login == 1 and st.session_state.page == 2:

    with st.expander("", expanded=True):
        st.subheader("InsightFox's Data Cleaner")
        st.write("Clean, Fill, and generally prepare your dataset for data processing and report creation.")

        uploaded_file = st.file_uploader("Upload your Excel or CSV File", type=["csv", "xlsx"])
        if uploaded_file is not None:
            try:
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                else:
                    df = pd.read_excel(uploaded_file)
                    
                if df.shape[0] > 10000 or df.shape[1] > 50:
                    st.error("The maximum is 5000 rows and 50 columns. Uploaded file was not saved.")
                    if 'df' in st.session_state:
                        del st.session_state['df'] #delete previous df
                    df = None # prevent the rest of the code from running

                else:
                    st.success("File uploaded successfully!")

            except Exception as e:
                st.error(f"Error reading file: {e}")
                df = None

            if df is not None:
                st.write("Uploaded Dataframe:")
                st.dataframe(df.head())

                st.session_state['df'] = df

    if uploaded_file is not None:
        if 'df' in st.session_state:
            df = st.session_state['df'].copy()

            # Column Selection
            columns = df.columns.tolist()

            if 'selected_columns' not in st.session_state:
                st.session_state['selected_columns'] = []

            with st.expander(" ", expanded=True):
                st.subheader("Select Columns to Analyze")
                st.write("""Click "All" for All Columns, and "Custom" for Custom Columns""")
                col_selection = st.radio("Select Columns:", ["All Columns", "Custom Columns"], key="col_selection")

                if col_selection == "All Columns":
                    selected_columns = columns
                    st.session_state['selected_columns'] = columns
                elif col_selection == "Custom Columns":
                    selected_columns = st.multiselect("Select Columns to Analyze:", columns, default=st.session_state['selected_columns'])
                    st.session_state['selected_columns'] = selected_columns
                else:
                    selected_columns = st.session_state['selected_columns']

            col1, col2, col3 = st.columns([1, 1, 1])
            #col3, col4 = st.columns([2,2])

            with col1 : 
                # Fill NaN Values
                with st.expander("", expanded=True):
                    st.subheader("Handle NaN (empty) Values")
                    fill_nan = st.checkbox("Yes", key="fill_nan")
                    if fill_nan:
                        st.write("""Click "All" for All Columns, and "Custom" for Custom Columns""")
                        st.write("""Choose to remove NaN (empty values) or fill them with AI Algorithm""")
                        nan_selection = st.radio("Select Columns for NaN Handling:", ["All", "Custom"], key="nan_selection")
                        if nan_selection == "All":
                            nan_columns = selected_columns
                        elif nan_selection == "Custom":
                            nan_columns = st.multiselect("Multiselect Variables", selected_columns, key="nan_columns")
                        else:
                            nan_columns = []
                        fill_or_remove = st.radio("Fill or Remove NaN", ["Fill", "Remove"], key="fill_or_remove")
            with col2 : 
                # Remove Duplicates
                with st.expander("", expanded=True):
                    st.subheader("Remove Duplicate Values")
                    remove_duplicates = st.checkbox("Yes", key="remove_duplicates")
                    if remove_duplicates:
                        st.write("""Click "All" for All Columns, and "Custom" for Custom Columns""")
                        duplicate_selection = st.radio("Select Columns for Duplicate Handling:", ["All", "Custom"], key="duplicate_selection")
                        if duplicate_selection == "All":
                            duplicate_columns = selected_columns
                        elif duplicate_selection == "Custom":
                            duplicate_columns = st.multiselect("Multiselect Variables", selected_columns, key="duplicate_columns")
                        else:
                            duplicate_columns = []
            with col3 : 
            # Remove Outliers
                with st.expander("", expanded=True):
                    st.subheader("Remove Outlier Values")
                    remove_outliers = st.checkbox("Yes", key="remove_outliers")
                    if remove_outliers:
                        st.write("""Click "All" for All Columns, and "Custom" for Custom Columns""")
                        outlier_selection = st.radio("Select Columns for Outlier Handling:", ["All", "Custom"], key="outlier_selection")
                        if outlier_selection == "All":
                            outlier_columns = selected_columns
                        elif outlier_selection == "Custom":
                            outlier_columns = st.multiselect("Multiselect Variables", selected_columns, key="outlier_columns")
                        else:
                            outlier_columns = []
            #with col4 : 
                # Choose Certain Category
                #with st.expander("Choose Certain Category", expanded=True):
                    #choose_category = st.checkbox("Yes", key="choose_category")
                    #if choose_category:
                        #category_column = st.selectbox("Select Column", selected_columns, key="category_column")
                        #unique_values = df[category_column].unique().tolist()
                        #selected_values = st.multiselect("Multiselect Strings Inside The Columns", unique_values, key="selected_values")

            # Check for Duplicates, NaNs, Outliers
            with st.expander("", expanded=True):
                st.subheader("Check for Duplicates/NaNs/Outliers")
                st.write("Check for Duplicates/NaNs/Outliers in your Current Data to better understand the problems inside the data")
                col5, col6, col7 = st.columns(3)
                if col5.button("Check Duplicates"):
                    duplicates = df[df.duplicated(subset=selected_columns, keep=False)]
                    st.write("Dataframe of Duplicates:")
                    st.dataframe(duplicates)

                if col6.button("Check NaNs"):
                    nan_rows = df[df.isnull().any(axis=1)]
                    nan_cols = df.loc[:, df.isnull().any(axis=0)]
                    st.write("Rows with NaN values:")
                    st.dataframe(nan_rows)
                    st.write("Columns with NaN values:")
                    st.dataframe(nan_cols)

                if col7.button("Check Outliers"):
                    outlier_rows = pd.DataFrame()
                    for col in selected_columns:
                        if pd.api.types.is_numeric_dtype(df[col]):
                            Q1 = df[col].quantile(0.25)
                            Q3 = df[col].quantile(0.75)
                            IQR = Q3 - Q1
                            lower_bound = Q1 - 1.5 * IQR
                            upper_bound = Q3 + 1.5 * IQR
                            outliers = df[(df[col] < lower_bound) | (df[col] > upper_bound)]
                            outlier_rows = pd.concat([outlier_rows, outliers])
                    st.write("Rows with Outliers:")
                    st.dataframe(outlier_rows.drop_duplicates())

        def remove_outliers_iqr(df, column):
            """Removes outliers from a DataFrame column using the IQR method."""
            Q1 = df[column].quantile(0.25)
            Q3 = df[column].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            df_filtered = df[(df[column] >= lower_bound) & (df[column] <= upper_bound)]
            return df_filtered

        if 'df' in st.session_state:
            df = st.session_state['df'].copy()

            st.markdown('<h3 style="font-size: 19px; text-align: left; color: #FFFFFF;">Click this to clean the data!!!</h3>', unsafe_allow_html=True)
            if st.button("Clean"):
                status_container_plots = st.empty()
                status_container_plots.info("Cleaning...")
                if fill_nan and nan_columns:
                    for col in nan_columns:
                        if fill_or_remove == "Fill":
                            if pd.api.types.is_numeric_dtype(df[col]):
                                df[col] = df[col].fillna(df[col].mean())
                            else:
                                df[col] = df[col].ffill()
                        elif fill_or_remove == "Remove":
                            df.dropna(subset=[col], inplace=True)

                if remove_duplicates and duplicate_columns:
                    df.drop_duplicates(subset=duplicate_columns, inplace=True)

                if remove_outliers and outlier_columns:
                    for col in outlier_columns:
                        if pd.api.types.is_numeric_dtype(df[col]):
                            df = remove_outliers_iqr(df, col)

                #if choose_category and selected_values:
                    #df = df[df[category_column].isin(selected_values)]

                #st.session_state['cleaned_df'] = df
                df_quota = get_quota_df()
                user_quota = df_quota[df_quota['email'] == st.session_state.email]
                if not user_quota.empty:
                        # quota_clean button
                        quota_clean_value = int(user_quota.iloc[0]["quota_clean"])
                        if quota_clean_value > 0:
                            update_quota(st.session_state.email, "quota_clean")
                            st.success(f"Success !! Quota left : {quota_clean_value - 1}")
                            st.success("Please wait for a few seconds, do not press anything until done")
                            st.session_state['cleaned_df'] = df
                            time.sleep(2)
                            st.rerun()  # Refresh to show updated quota
                        else:
                            st.error("No quota left!")

        if 'cleaned_df' in st.session_state :
            cleaned_df = st.session_state['cleaned_df']

            def to_excel(df):
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Sheet1')
                processed_data = output.getvalue()
                return processed_data

            excel_file = to_excel(cleaned_df)
            st.download_button(
                label="Download Cleaned Data as Excel",
                data=excel_file,
                file_name="cleaned_data.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

# ============================================== EDA REPORTING ==========================================================================================

if st.session_state.login == 1 and st.session_state.page == 3:

    all_plot_options = ["Histograms", "Boxplots", "Lineplots", "Areaplots",
                         "Violinplots", "Correlation Heatmap", "CDF", "Categorical Barplots",
                         "Piecharts", "Wordclouds", "Countplots", "Treemaps"]

    with st.expander("", expanded=True):
        st.subheader("InsightFox's Exploratory Data Analysis Reports and Data Plotting")
        st.write("Generate Exploratory Data Analysis Reports and Plot All of the Data Analytics of your data")

        uploaded_file = st.file_uploader("Upload your Excel or CSV File", type=["csv", "xlsx"], accept_multiple_files=False)

        if uploaded_file is not None:
            try:
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                else:
                    df = pd.read_excel(uploaded_file)

                if df.shape[0] > 100000 or df.shape[1] > 50:
                    st.error("The maximum is 5000 rows and 50 columns. Uploaded file was not saved.")
                    if 'df' in st.session_state:
                        del st.session_state['df'] #delete previous df
                    df = None  # Prevent further processing
                else:
                    st.success("File uploaded successfully!")

            except Exception as e:
                st.error(f"Error reading file: {e}")
                df = None

            if df is not None:
                st.write("Uploaded Dataframe:")
                st.dataframe(df.head())

                # New code: Let the user select the columns to analyze
                columns_to_analyze = st.multiselect(
                    "Select Columns to Analyze",
                    options=df.columns.tolist(),
                    default=df.columns.tolist(),
                    help="Choose the columns that you want to include in the analysis."
                )

                # Update the session state with the filtered DataFrame
                st.session_state['df'] = df[columns_to_analyze]

    if 'plot_fig' not in st.session_state:
        st.session_state['plot_fig'] = None

    if uploaded_file is not None:
        if 'df' in st.session_state:
            col1, col2, col3 = st.columns([2, 2, 2])

            with col1:
                # Plot Examples (Expander)
                with st.expander("", expanded=True):
                    st.subheader("Plot Examples")
                    st.write("Confused of what plots to generate? see these examples for yourselves!")

                    example_plot_choice = st.selectbox("Choose Plot", options=all_plot_options, index=0)

                    theme_choice = st.selectbox(
                        "Select Plot Theme",
                        ["Orange", "Red", "Green", "Purple", "Blue", "Gray", "Pastel"],
                        index=0
                    )

                    if example_plot_choice == "Histograms":
                        st.session_state['plot_fig'] = plot_histograms(df_sample_int, theme=theme_choice)

                    elif example_plot_choice == "Boxplots":
                        st.session_state['plot_fig'] = plot_boxplots(df_sample_int, theme=theme_choice)

                    elif example_plot_choice == "Lineplots":
                        st.session_state['plot_fig'] = plot_lineplots(df_sample_int, theme=theme_choice)

                    elif example_plot_choice == "Areaplots":
                        st.session_state['plot_fig'] = plot_areaplots(df_sample_int, theme=theme_choice)

                    elif example_plot_choice == "Violinplots":
                        st.session_state['plot_fig'] = plot_violinplots(df_sample_int, theme=theme_choice)

                    elif example_plot_choice == "Correlation Heatmap":
                        st.session_state['plot_fig'] = plot_correlation_heatmap(df_sample_int, cmap_option=theme_choice)

                    elif example_plot_choice == "CDF":
                        st.session_state['plot_fig'] = plot_cdf(df_sample_int, theme=theme_choice)

                    elif example_plot_choice == "Categorical Barplots":
                        st.session_state['plot_fig'] = plot_categorical_barplots(df_sample_str, theme=theme_choice)

                    elif example_plot_choice == "Piecharts":
                        st.session_state['plot_fig'] = plot_piecharts(df_sample_str, theme=theme_choice)

                    elif example_plot_choice == "Stacked Barplots":
                        st.session_state['plot_fig'] = plot_stacked_barplots(df_sample_str, theme=theme_choice)

                    elif example_plot_choice == "Grouped Barplots":
                        st.session_state['plot_fig'] = plot_grouped_barplots(df_sample_str, theme=theme_choice)

                    elif example_plot_choice == "Wordclouds":
                        st.session_state['plot_fig'] = plot_wordclouds(df_sample_str, theme=theme_choice)

                    elif example_plot_choice == "Countplots":
                        st.session_state['plot_fig'] = plot_countplots(df_sample_str, theme=theme_choice)

                    elif example_plot_choice == "Treemaps":
                        st.session_state['plot_fig'] = plot_treemaps(df_sample_str, theme=theme_choice)

                if st.session_state['plot_fig'] is not None:
                    st.pyplot(st.session_state['plot_fig'], clear_figure=True)

            with col2:
                with st.expander("", expanded=True):
                    st.subheader("Key Statistics Generator")
                    st.markdown("Generate the key statistics of your Excel or CSV File and creates a :orange[Ms. Word Report]")
                    st.warning("Please do not run two buttons at once, wait until one service is done before initiating another")
                    if st.button("Generate Key Statistics"):
                        status_container = st.empty()
                        status_container.info("Generating Key Statistics reports...")
                        st.error("Please do not press any other button while the report is being generated!!!")
                        df_quota = get_quota_df()
                        user_quota = df_quota[df_quota['email'] == st.session_state.email]
                        if not user_quota.empty:
                            quota_plot_value = int(user_quota.iloc[0]["quota_plot"])
                            if quota_plot_value > 0:
                                update_quota(st.session_state.email, "quota_plot")
                                st.success(f"Success !! Quota left : {quota_plot_value - 1}")
                                st.success("Please wait for a few seconds, we're initiating the download button...")

                                # Determine language based on session state (or other logic)
                                language = st.session_state.get('language', 'en')  # Default to English if not set

                                buffer_basic = BytesIO()
                                buffer_columnwise = BytesIO()

                                doc_basic_en = eda_dataframe_to_docx_en(st.session_state['df'])
                                doc_columnwise_en = eda_all_columns_to_docx_en(st.session_state['df'])
                                doc_basic_id = eda_dataframe_to_docx_id(st.session_state['df'])
                                doc_columnwise_id = eda_all_columns_to_docx_id(st.session_state['df'])

                                doc_basic_en.save(buffer_basic)
                                buffer_basic.seek(0)
                                doc_columnwise_en.save(buffer_columnwise)
                                buffer_columnwise.seek(0)

                                buffer_basic_id = BytesIO()
                                buffer_columnwise_id = BytesIO()

                                doc_basic_id.save(buffer_basic_id)
                                buffer_basic_id.seek(0)
                                doc_columnwise_id.save(buffer_columnwise_id)
                                buffer_columnwise_id.seek(0)

                                zip_buffer = BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                                    zipf.writestr("EN - Basic Key Statistics Report.docx", buffer_basic.getvalue())
                                    zipf.writestr("EN - Columnwise Key Statistics Report.docx", buffer_columnwise.getvalue())
                                    zipf.writestr("ID - Basic Key Statistics Report.docx", buffer_basic_id.getvalue())
                                    zipf.writestr("ID - Columnwise Key Statistics Report.docx", buffer_columnwise_id.getvalue())
                                zip_buffer.seek(0)

                                status_container.success("Key Statistics reports generated!")
                                st.session_state['eda_zip_buffer'] = zip_buffer
                                time.sleep(6)
                            else:
                                st.error("No quota left!")

                    if 'eda_zip_buffer' in st.session_state:
                        st.download_button(
                            label="Download Key Statistics Reports (ZIP)",
                            data=st.session_state['eda_zip_buffer'],
                            file_name="ID_EN Key Statistics Report.zip",
                            mime="application/zip"
                        )

                with st.expander("", expanded=True):
                    st.subheader("Exploratory Data Analysis")
                    st.markdown("Generate the key statistics and analyzes (paragraphed) of your Excel or CSV File and creates a :orange[Ms. Word Report]")
                    st.warning("Please do not run two buttons at once, wait until one service is done before initiating another")
                    timeseries = st.checkbox("Data is Time Series")
                    if timeseries : 
                        ts_choice = st.selectbox("Select the datetime column for time series analysis", options=df.columns.tolist())
                        #df.set_index(ts_choice, inplace=True)
                        timefreq = st.selectbox(
                            "Select the frequency of the data",
                            options=["Seconds", "Minutes", "Hours", "Days", "Weeks", "Months", "Quarterly", "Yearly"])
                    if st.button("Generate Report"):
                        status_container = st.empty()
                        status_container.info("Generating Exploratory Data Analysis reports...")
                        st.error("Please do not press any other button while the report is being generated!!!")
                        df_quota = get_quota_df()
                        user_quota = df_quota[df_quota['email'] == st.session_state.email]
                        if not user_quota.empty:
                            quota_plot_value = int(user_quota.iloc[0]["quota_plot"])
                            if quota_plot_value > 0:
                                update_quota(st.session_state.email, "quota_plot")
                                st.success(f"Success !! Quota left : {quota_plot_value - 1}")
                                st.success("Please wait for a few seconds, we're initiating the download button...")
                                if timeseries : 
                                    doc_en = generate_doc_report_en_ts(st.session_state['df'])
                                else : 
                                    doc_en = generate_doc_report_en(st.session_state['df'])
                                buffer_en = BytesIO()
                                doc_en.save(buffer_en)
                                buffer_en.seek(0)
                                
                                if timeseries :
                                    doc_id = generate_doc_report_id_ts(st.session_state['df'])
                                else : 
                                    doc_id = generate_doc_report_id(st.session_state['df'])
                                buffer_id = BytesIO()
                                doc_id.save(buffer_id)
                                buffer_id.seek(0)

                                zip_doc = BytesIO()
                                with zipfile.ZipFile(zip_doc, 'w') as zipd:
                                    zipd.writestr("ENGLISH - Key Insights Analysis.docx", buffer_en.getvalue())
                                    zipd.writestr("INDONESIAN - Key Insights Analysis.docx", buffer_id.getvalue())
                                zip_doc.seek(0)

                                #status_container.success("Key Insight reports generated!")
                                st.session_state['eda_zip_doc'] = zip_doc
                                time.sleep(6)
                            else:
                                st.error("No quota left!")

                    if 'eda_zip_doc' in st.session_state:
                        st.download_button(
                            label="Download EDA Reports (ZIP)",
                            data=st.session_state['eda_zip_doc'],
                            file_name="ID_EN Exploratory Data Analysis Reports.zip",
                            mime="application/zip"
                        )

            with col3:
                with st.expander("", expanded=True):
                    st.subheader("Finished Plot Downloader")
                    st.markdown("Generate the desired plots (graphs, bars, etc) of your Excel or CSV File. Downloads a :orange[ZIP File containing all the plots (Max : 5 Types).]")
                    st.warning("Please do not run two buttons at once, wait until one service is done before initiating another")
                    selected_plot_types = st.multiselect(
                        "Choose plots to download",
                        options=[
                            "Histograms", "Boxplots", "Scatterplots", "Lineplots", "Areaplots",
                            "Violinplots", "Correlation Heatmap", "CDF", "Categorical Barplots",
                            "Piecharts", "Stacked Barplots", "Grouped Barplots", "Wordclouds",
                            "Countplots", "Treemaps"
                        ],
                        default=[
                            "Histograms", "Areaplots", 
                            "Correlation Heatmap", "Categorical Barplots",
                        ],
                        max_selections=5
                    )
                    theme = st.selectbox("Select Theme", ["Orange", "Green", "Red", "Purple", "Blue", "Gray", "Pastel"], index=0)
                    if st.button("Generate Visualizations"):
                        status_container_plots = st.empty()
                        status_container_plots.info("Generating plots...")
                        st.error("Please do not press any other button while the plots are being generated!!!")
                        df_quota = get_quota_df()
                        user_quota = df_quota[df_quota['email'] == st.session_state.email]
                        if not user_quota.empty:
                            quota_plot_value = int(user_quota.iloc[0]["quota_plot"])
                            if quota_plot_value > 0:
                                update_quota(st.session_state.email, "quota_plot")
                                zip_buffer_plots = aggregate_and_download_plots(st.session_state['df'], selected_plot_types, theme)
                                #status_container_plots.success("Plots generated!")
                                st.session_state['plot_zip_buffer'] = zip_buffer_plots
                                st.success(f"Success !! Quota left : {quota_plot_value - 1}")
                                st.success("Please wait for a few seconds, we're initiating the download button...")
                                time.sleep(6)
                            else:
                                st.error("No quota left!")

                    if 'plot_zip_buffer' in st.session_state:
                        st.download_button(
                            label="Download Selected Visualizations (ZIP)",
                            data=st.session_state['plot_zip_buffer'],
                            file_name="Data Visualizations.zip",
                            mime="application/zip"
                        )

# ========================================== BUYING QUOTA ======================================================================================

if st.session_state.login == 1 and st.session_state.page == 4:
    
    # Load apikey dataframe if not already in session state
    if 'df_apikey' not in st.session_state:
        apisheetskey = "1sIEI-_9N96ndRJgWDyl0iL65bACeGQ74MncOV4HQCXY"
        url_apikey = f'https://docs.google.com/spreadsheet/ccc?key={apisheetskey}&output=csv'
        st.session_state.df_apikey = pd.read_csv(url_apikey)
    
    df_apikey = st.session_state.df_apikey
    platform = "Gemini"
    apikeyxloc = df_apikey['Platform'].str.contains(platform).idxmax()
    apikey = df_apikey.iloc[apikeyxloc, 2]
    
    # Configure genai only once
    if 'genai_configured' not in st.session_state:
        genai.configure(api_key=apikey)
        st.session_state.genai_configured = True
    
    # Load quota dataframe if not already in session state
    if 'df_quota' not in st.session_state:
        url = f'https://docs.google.com/spreadsheet/ccc?key={sheetskey}&output=xlsx'
        st.session_state.df_quota = pd.read_excel(url, sheet_name="Quotas")
    
    df_quota = st.session_state.df_quota
    user_row = df_quota[df_quota['email'] == st.session_state.email]
    if not user_row.empty:
        quota_clean = user_row['quota_clean'].values[0]
        quota_plot = user_row['quota_plot'].values[0]
        with st.expander("", expanded=True) : 
            st.markdown('<h3 style="font-size: 35px; text-align: center; color: #28282b;">MY QUOTA</h3>', unsafe_allow_html=True)
            st.subheader(f"Account In Use : {st.session_state.email}")
            st.success(f"Cleaning Quota: {quota_clean}")
            st.success(f"EDA Report, Key Stats Report, and Data Visualization Quota : {quota_plot}")
    else:
        st.warning("User not found in quota records.")

    st.subheader("")
    st.markdown('<h3 style="font-size: 35px; text-align: center; color: #FFFFFF;">BUY QUOTAS</h3>', unsafe_allow_html=True)

    colbuy1, colbuy2, colbuy3, colbuy4 = st.columns([1,1,1,1])

    with colbuy1: 
        with st.expander("", expanded=True) : 
            st.markdown('<h3 style="font-size: 35px; text-align: center; color: #28282b;">smol package</h3>', unsafe_allow_html=True)
            st.write("")
            st.markdown("starterpack... you can automate 2 reports for 2 files, or a library of an immense amount plots for :orange[2 files] with the maximum being :orange[50 columns] and :orange[5000 rows] per file.... yk 15k is kinda rlly cheap for all of this")
            st.markdown("This means that one report/A WHOLE LOT of plots is :orange[about IDR 7k]")
            st.write("")
            st.success("Cleaning Quota : 1")
            st.success("EDA Report, Key Stats Report, and Data Visualization Quota : 2")
            st.write("")
            st.markdown('<div style="background-color: #bd5c34; padding: 3px; border-radius: 15px; width: 80%; height: 50%; margin: auto; display: flex; justify-content: center; align-items: center;"><h1 style="color: #FFFFFF; font-size: 24px; margin-left: 22px;">15k</h1></div>', unsafe_allow_html=True)
            st.write("")
            st.write("")

    with colbuy2: 
        with st.expander("", expanded=True) : 
            st.markdown('<h3 style="font-size: 35px; text-align: center; color: #28282b;">skibidi package</h3>', unsafe_allow_html=True)
            st.write("")
            st.markdown("This is cool ig... you can automate 4 reports for 4 files, or a library of an immense amount plots for :orange[4 files] with the maximum being :orange[50 columns] and :orange[5000 rows] per file.... yk 35k is kinda rlly cheap for all of this")
            st.markdown("This means that one report/A WHOLE LOT of plots is :orange[about IDR 7k]")
            st.write("")
            st.success("Cleaning Quota : 2")
            st.success("EDA Report, Key Stats Report, and Data Visualization Quota : 4")
            st.write("")
            st.markdown('<div style="background-color: #bd5c34; padding: 3px; border-radius: 15px; width: 80%; height: 50%; margin: auto; display: flex; justify-content: center; align-items: center;"><h1 style="color: #FFFFFF; font-size: 24px; margin-left: 22px;">35k</h1></div>', unsafe_allow_html=True)
            st.write("")
            st.write("")

    with colbuy3: 
        with st.expander("", expanded=True) : 
            st.markdown('<h3 style="font-size: 35px; text-align: center; color: #28282b;">Rizz Package</h3>', unsafe_allow_html=True)
            st.write("")
            st.markdown("It gets better?! now you have double the quota for :orange[less than the price of two...] you can automate the reports  or create a library of an immense amount plots for of :orange[8 files] containing a maximum of :orange[50 columns] and :orange[5000 rows] each....")
            st.markdown("This means that one report/A WHOLE LOT of plots is :orange[about IDR 7k]")
            st.write("")
            st.success("Cleaning Quota : 4")
            st.success("EDA Report, Key Stats Report, and Data Visualization Quota : 8")
            st.write("")
            st.markdown('<div style="background-color: #bd5c34; padding: 3px; border-radius: 15px; width: 80%; height: 50%; margin: auto; display: flex; justify-content: center; align-items: center;"><h1 style="color: #FFFFFF; font-size: 24px; margin-left: 22px;">60k</h1></div>', unsafe_allow_html=True)
            st.write("")
            st.write("")

    with colbuy4: 
        with st.expander("", expanded=True) : 
            st.markdown('<h3 style="font-size: 35px; text-align: center; color: #28282b;">SIGMA PACKAGE</h3>', unsafe_allow_html=True)
            st.write("")
            st.markdown("Now this is GAS!! :fire::fire::fire: now you have TRIPLE the quota for :orange[significantly less than the price of three...], automate the reports  or a create library of an immense amount plots for of :orange[12 files] with a max :orange[50 columns] and :orange[5000 rows] each....")
            st.markdown("This means that one report/A WHOLE LOT of plots is :orange[LESS THAN IDR 7k]")
            st.write("")
            st.success("Cleaning Quota : 6")
            st.success("EDA Report, Key Stats Report, and Data Visualization Quota : 12")
            st.write("")
            st.markdown('<div style="background-color: #bd5c34; padding: 3px; border-radius: 15px; width: 80%; height: 50%; margin: auto; display: flex; justify-content: center; align-items: center;"><h1 style="color: #FFFFFF; font-size: 24px; margin-left: 22px;">85k</h1></div>', unsafe_allow_html=True)
            st.write("")
            st.write("")  

    with st.expander("", expanded=True) : 
        st.header("Self-Checkout")
        st.write("Proceed Payment through QRIS provided and upload the transaction proof")
        st.warning("Make sure to only pay these EXACT amounts : IDR 15.000, IDR 35.000, IDR 60.000, IDR 85.000. The system will NOT detect other nominal payments.")
        st.write("")
        colqris1, colqris2, colqris3 = st.columns([1, 1, 1])
        with colqris2:
            st.image("QRIS_Keno.jpg", use_container_width=True)
        uploaded_file = st.file_uploader("Upload Transaction Image (JPEG or PNG)", type=["jpg", "jpeg", "png"])

        if uploaded_file is not None:
            image_parts = input_image_setup(uploaded_file)

            input_prompt = """
            Please extract the amount of money (along with its currency) that was transferred in the transaction along with the time of transaction.
            Just return me the currency, amount, and time. Follow the example below strictly. Use "." dot and not comma for the numbers.
            Remember, JUST INPUT LIKE THE EXAMPLE, NO NEED ANYTHING ELSE, NOT EVEN AN INTRO. 

            Example : 
            IDR - 45.000 : 19:21 (HH-MM)
            IDR - 35.000 : 21:07 (HH-MM)
            """

            if st.button("Verify Payment"):
                with st.spinner("Processing..."):
                    response_text = get_gemini_repsonse(input_prompt, image_parts)
                    print("")
                    print("")
                    print(response_text)
                    print("")
                    print("")
                    try:
                        time_str = response_text.split(":")[-1].strip()  # Extract time part
                        transaction_time = time_str.split(" ")[0].strip()  # Remove "(HH-MM)"
                        image_hash = generate_image_hash(transaction_time)  # Hash the time string

                        if check_image_hash_exists(st.session_state.email, image_hash):
                            st.error("You have already paid with this transaction!")
                        else:
                            if "15.000" in response_text:
                                add_quota(st.session_state.email, "quota_clean", 1)
                                add_quota(st.session_state.email, "quota_plot", 2)
                                st.success("Payment Successful!!!")
                                st.warning("Please re-login (refresh the page) to refresh the quota")
                            if "35.000" in response_text:
                                add_quota(st.session_state.email, "quota_clean", 2)
                                add_quota(st.session_state.email, "quota_plot", 4)
                                st.success("Payment Successful!!!")
                                st.warning("Please re-login (refresh the page) to refresh the quota")
                            elif "60.000" in response_text:
                                add_quota(st.session_state.email, "quota_clean", amount_to_add=4)
                                add_quota(st.session_state.email, "quota_plot", amount_to_add=8)
                                st.success("Payment Successful!!!")
                                st.warning("Please re-login (refresh the page) to refresh the quota")
                            elif "85.000" in response_text:
                                add_quota(st.session_state.email, "quota_clean", amount_to_add=6)
                                add_quota(st.session_state.email, "quota_plot", amount_to_add=12)
                                st.success("Payment Successful!!!")
                                st.warning("Please re-login (refresh the page) to refresh the quota")
                            else:
                                st.error("Self checkout failed.")
                                st.error("Please refer to contacting insightfoxa@gmail.com so we can manually add your quota, Include Transaction Proof and Account wish to be granted quota.")

                            update_image_hashes(st.session_state.email, image_hash)
                    except:
                        st.error("Self checkout failed.")
                        st.error("Please refer to contacting insightfoxa@gmail.com so we can manually add your quota, Include Transaction Proof and Account wish to be granted quota.")
